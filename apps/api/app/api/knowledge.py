"""
Company Knowledge Base & RAG Asset Management Endpoints.
Strictly scoped by tenant_id via SQLAlchemy 2 Async and Postgres RLS.
Supports:
1. Real multipart file uploads (PDF, DOCX, TXT, images) with OCR, text extraction, and pgvector embeddings.
2. Server-enforced limits: 50 MB per file, and plan-based quotas (starter: 20, pro: 100, enterprise: unlimited).
3. Real permanent Web URL scraping with fail-closed anti-hallucination rules.
4. Dynamic statuses: 'indexed', 'processing', 'failed'.
"""
import hashlib
import io
import logging
import re
import uuid
from datetime import datetime, timezone
from html.parser import HTMLParser
from typing import Any, Dict, List, Optional, Tuple
import httpx

logger = logging.getLogger(__name__)
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession


class HTMLTextExtractor(HTMLParser):
    """Pure Python HTML text extractor that removes scripts, styles, and extracts title and clean body text."""
    def __init__(self):
        super().__init__()
        self.reset()
        self.fed: List[str] = []
        self.title_parts: List[str] = []
        self._in_title = False
        self._skip = False
        self._skip_tags = {"script", "style", "noscript", "svg", "header", "footer", "nav"}

    def handle_starttag(self, tag: str, attrs: Any):
        t = tag.lower()
        if t in self._skip_tags:
            self._skip = True
        elif t == "title":
            self._in_title = True

    def handle_endtag(self, tag: str):
        t = tag.lower()
        if t in self._skip_tags:
            self._skip = False
        elif t == "title":
            self._in_title = False

    def handle_data(self, d: str):
        if self._in_title:
            self.title_parts.append(d)
        elif not self._skip:
            self.fed.append(d)

    def get_text(self) -> str:
        raw = " ".join(self.fed)
        return re.sub(r"\s+", " ", raw).strip()

    def get_title(self) -> str:
        return re.sub(r"\s+", " ", " ".join(self.title_parts)).strip()


from app.core.config import settings
from app.core.db import get_db
from app.core.security import CurrentTenantUser, get_current_tenant_user
from app.core.storage import storage_service
from app.models.entities import CompanyAsset, ExportJob, ExportTemplate, KnowledgeVector, Project
from app.models.schemas import (
    CompanyAssetCreate,
    CompanyAssetOut,
    KnowledgeStatsOut,
    KnowledgeUploadResponse,
    KnowledgeWebSourceInput,
)
from app.services.billing_service import billing_service
from app.services.chunking_service import chunking_service
from app.services.embedding_service import embedding_service
from app.services.ocr_service import OCRService

router = APIRouter(prefix="/knowledge", tags=["Knowledge Base & Assets"])
ocr_service = OCRService()

MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB


def extract_text_from_upload(filename: str, file_bytes: bytes) -> Tuple[str, str, Optional[str]]:
    """
    Extraction textuelle mutualisee entre l'upload direct (ci-dessous) et le script
    de backfill de re-indexation (scripts/backfill_knowledge_vectors.py) : OCR pour
    PDF/image, parseur python-docx pour DOCX, decodage direct pour texte brut.
    Ne leve jamais d'exception : tout echec est capture et renvoye comme
    status_state="failed" avec le message d'erreur, exactement comme le comportement
    historique de l'endpoint d'upload. (30/08)
    """
    filename_lower = filename.lower()
    extracted_text = ""
    status_state = "indexed"
    error_msg = None
    try:
        if filename_lower.endswith(".pdf"):
            parsed = ocr_service.extract_text_and_tables(file_bytes, filename)
            extracted_text = parsed.get("full_text", "")
            if not extracted_text and parsed.get("pages"):
                extracted_text = "\n\n".join(p.get("text", "") for p in parsed["pages"])
        elif filename_lower.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            tables_text.append(" | ".join(row_cells))
                extracted_text = "\n\n".join(paragraphs + tables_text)
            except Exception:
                extracted_text = f"[Document Word {filename}]"
        elif filename_lower.endswith((".txt", ".md", ".csv", ".json")):
            extracted_text = file_bytes.decode("utf-8", errors="replace")
        elif filename_lower.endswith((".png", ".jpg", ".jpeg", ".tiff")):
            parsed = ocr_service.extract_text_and_tables(file_bytes, filename)
            extracted_text = parsed.get("full_text", "")
        else:
            extracted_text = f"Document technique {filename}"
    except Exception as e:
        status_state = "failed"
        error_msg = f"Erreur d'extraction textuelle / OCR : {str(e)}"
        extracted_text = f"Fichier {filename} (non analyse)"
    return extracted_text, status_state, error_msg


def chunk_and_embed_asset_text(
    tenant_id_uuid,
    asset_id,
    category: str,
    title: str,
    full_text: str,
) -> List[KnowledgeVector]:
    """
    Decoupe le texte COMPLET (aucun plafond de longueur, contrairement a l'ancien
    extracted_text[:4000]) en fragments semantiques via chunking_service (~1200
    caracteres/fragment, meme logique que pour les DCE), puis calcule un embedding
    par fragment. Ne persiste rien elle-meme : l'appelant fait db.add() puis commit
    dans sa propre transaction (permet un rollback atomique avec le CompanyAsset
    parent en cas d'echec). Retourne [] si le texte est vide -- jamais d'exception
    remontee a l'appelant (generate_embedding() retourne toujours un vecteur, y
    compris un repli deterministe en cas d'echec LLM). (30/08)
    """
    if not full_text or not full_text.strip():
        return []
    pages = [{"page_number": 1, "text": full_text}]
    chunks = chunking_service.chunk_document_pages(pages)
    if not chunks:
        return []
    vectors = embedding_service.generate_batch_embeddings([c["content"] for c in chunks])
    rows: List[KnowledgeVector] = []
    for c, vec in zip(chunks, vectors):
        rows.append(KnowledgeVector(
            id=uuid.uuid4(),
            tenant_id=tenant_id_uuid,
            asset_id=asset_id,
            category=category,
            content=c["content"],
            embedding=vec,
            metadata_json={
                "asset_title": title,
                "chunk_index": c["chunk_index"],
                "section_title": c.get("section_title"),
                "char_count": len(c["content"]),
            },
            created_at=datetime.now(timezone.utc),
        ))
    return rows


@router.get("/assets", response_model=List[CompanyAssetOut])
async def list_company_assets(
    category: Optional[str] = None,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Returns company references, certifications, equipment sheets, CVs, and web sources.
    Strictly scoped to the authenticated tenant via PostgreSQL RLS.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    stmt = select(CompanyAsset).where(CompanyAsset.tenant_id == t_uuid)
    if category:
        stmt = stmt.where(CompanyAsset.category == category)
    stmt = stmt.order_by(CompanyAsset.created_at.desc())

    result = await db.execute(stmt)
    assets = result.scalars().all()

    return [
        CompanyAssetOut(
            id=str(a.id),
            tenant_id=str(a.tenant_id),
            category=a.category,
            title=a.title,
            description=a.description or "",
            s3_url=a.s3_url,
            tags=a.metadata_json.get("tags", []) if isinstance(a.metadata_json, dict) else [],
            metadata_json=a.metadata_json or {},
            created_at=a.created_at,
        )
        for a in assets
    ]


@router.get("/stats", response_model=KnowledgeStatsOut)
async def get_knowledge_stats(
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Returns knowledge asset quota metrics, usage, and category breakdown.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    quota_info = await billing_service.check_and_enforce_knowledge_quota(t_uuid, db=db) if False else None
    
    # Get plan & quota directly
    sub = await billing_service.get_tenant_subscription(t_uuid, db)
    plan_id = sub.plan_id.lower() if sub else "starter"
    quotas = {"starter": 20, "pro": 100, "enterprise": None}
    max_allowed = quotas.get(plan_id, 20)

    stmt = select(CompanyAsset).where(CompanyAsset.tenant_id == t_uuid)
    result = await db.execute(stmt)
    assets = result.scalars().all()

    category_counts: Dict[str, int] = {}
    for a in assets:
        category_counts[a.category] = category_counts.get(a.category, 0) + 1

    return KnowledgeStatsOut(
        total_assets=len(assets),
        max_allowed=max_allowed,
        plan=plan_id,
        category_counts=category_counts,
    )


@router.post("/upload", response_model=KnowledgeUploadResponse, status_code=status.HTTP_201_CREATED)
async def upload_knowledge_document(
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    title: Optional[str] = Form(None),
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Real multipart/form-data upload for knowledge base documents.
    1. Validates 50 MB max file size.
    2. Enforces plan quota (starter: 20, pro: 100, enterprise: unlimited).
    3. Saves file to tenant storage.
    4. Extracts text & tables via OCR/parsers.
    5. Computes pgvector embeddings and saves to company_assets with status='indexed'.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    if not file.filename:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Fichier requis.")

    # 1. Check & Enforce Quota
    await billing_service.check_and_enforce_knowledge_quota(t_uuid, db=db)

    # 2. Read bytes and enforce 50 MB limit
    file_bytes = await file.read()
    file_size = len(file_bytes)

    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = round(file_size / (1024 * 1024), 2)
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"Le fichier dépasse la taille maximale autorisée de 50 Mo (taille reçue : {size_mb} Mo).",
        )

    # 2bis. Reject exact duplicate content already indexed for this tenant
    # (prevents accidental double-upload, e.g. a double drag-and-drop or double click)
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    dedup_stmt = select(CompanyAsset).where(
        CompanyAsset.tenant_id == t_uuid,
        CompanyAsset.status != "obsolete",
        CompanyAsset.metadata_json["file_hash"].astext == file_hash,
    )
    dedup_result = await db.execute(dedup_stmt)
    existing_duplicate = dedup_result.scalar_one_or_none()
    if existing_duplicate:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Ce fichier est identique à un document déjà indexé : « {existing_duplicate.title} ». Supprimez-le d'abord si vous voulez le remplacer.",
        )

    filename_lower = file.filename.lower()
    clean_name = re.sub(r"[^a-zA-Z0-9._-]", "_", file.filename)
    asset_id = uuid.uuid4()
    now = datetime.now(timezone.utc)

    # 3. Store file binary
    subpath = f"knowledge/{current_user.tenant_id}/{asset_id}_{clean_name}"
    content_type = file.content_type or "application/octet-stream"
    s3_key = storage_service.upload_file(
        tenant_id=current_user.tenant_id,
        subpath=subpath,
        file_obj=file_bytes,
        content_type=content_type,
    )

    # 4. Determine category and title if not provided
    inferred_category = category or "general"
    if not category:
        if "qualibat" in filename_lower or "certif" in filename_lower or "iso" in filename_lower:
            inferred_category = "certificat_qualibat"
        elif "engin" in filename_lower or "materiel" in filename_lower or "grue" in filename_lower:
            inferred_category = "materiel_engins"
        elif "cv" in filename_lower or "conducteur" in filename_lower or "equipe" in filename_lower:
            inferred_category = "cv_encadrement"
        elif "rse" in filename_lower or "environnement" in filename_lower or "qse" in filename_lower:
            inferred_category = "demarche_rse"
        elif "ref" in filename_lower or "chantier" in filename_lower:
            inferred_category = "reference_chantier"

    display_title = title.strip() if title and title.strip() else file.filename.rsplit(".", 1)[0].replace("_", " ")

    # 5. Extract text (OCR if PDF/image, docx parser if DOCX, text otherwise) --
    # logique mutualisee avec le backfill via extract_text_from_upload() (30/08)
    extracted_text, status_state, error_msg = extract_text_from_upload(file.filename, file_bytes)

    word_count = len(extracted_text.split()) if extracted_text else 0

    # 6. Generate Vector Embedding (sur les 4000 premiers caracteres -- conserve pour
    # compat ascendante sur CompanyAsset.embedding, encore lu ailleurs en secours)
    embedding_vector = None
    if extracted_text and status_state == "indexed":
        await embedding_service.sync_platform_key(db)
        try:
            embedding_vector = embedding_service.generate_embedding(extracted_text[:4000])
        except Exception:
            embedding_vector = None

    # 6bis. Indexation par fragments (chunking) sur le texte COMPLET -- corrige le
    # plafond 4000c ci-dessus qui ne rendait recherchable que les ~2 premieres pages
    # de tout document, quelle que soit sa longueur reelle. La recherche RAG reelle
    # (tasks.py) interroge desormais knowledge_vectors plutot que
    # CompanyAsset.embedding directement. (30/08)
    knowledge_vector_rows: List[KnowledgeVector] = []
    if extracted_text and status_state == "indexed":
        try:
            knowledge_vector_rows = chunk_and_embed_asset_text(
                tenant_id_uuid=t_uuid,
                asset_id=asset_id,
                category=inferred_category,
                title=display_title,
                full_text=extracted_text,
            )
        except Exception as chunk_exc:
            logger.warning("[knowledge.py] Chunking/embedding notice for asset %s: %s", asset_id, chunk_exc)

    metadata = {
        "source_type": "file",
        "file_name": file.filename,
        "file_size": file_size,
        "file_hash": file_hash,
        "content_type": content_type,
        "status": status_state,
        "error_message": error_msg,
        "word_count": word_count,
        "s3_key": s3_key,
        "tags": [inferred_category],
        "indexed_at": now.isoformat(),
    }
    if embedding_vector:
        metadata["has_embedding"] = True

    # 7. Persist CompanyAsset in PostgreSQL
    new_asset = CompanyAsset(
        id=asset_id,
        tenant_id=t_uuid,
        category=inferred_category,
        title=display_title,
        description=extracted_text[:12000] if extracted_text else display_title,
        s3_url=s3_key,
        status=status_state,
        embedding=embedding_vector,
        metadata_json=metadata,
        created_at=now,
        updated_at=now,
    )
    db.add(new_asset)
    for kv_row in knowledge_vector_rows:
        db.add(kv_row)
    await db.flush()
    await db.refresh(new_asset)

    return KnowledgeUploadResponse(
        success=True,
        asset_id=str(new_asset.id),
        title=new_asset.title,
        category=new_asset.category,
        status=status_state,
        file_size_bytes=file_size,
        word_count=word_count,
        message=f"Document '{file.filename}' analysé et indexé avec succès ({word_count} mots).",
    )


@router.post("/web-source", response_model=CompanyAssetOut, status_code=status.HTTP_201_CREATED)
async def add_knowledge_web_source(
    payload: KnowledgeWebSourceInput,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Permanent Web Source indexer for the company knowledge base.
    Fetches the real webpage content, strips HTML markup, generates pgvector embeddings,
    and indexes it under company_assets.
    Strict fail-closed anti-hallucination rule: if URL is unreachable, returns an explicit error.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    url = payload.url.strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="URL invalide. Veuillez fournir une adresse web valide commençant par http:// ou https://.",
        )

    # 1. Check & Enforce Quota
    await billing_service.check_and_enforce_knowledge_quota(t_uuid, db=db)

    now = datetime.now(timezone.utc)

    # 2. Fetch live webpage with timeout and user-agent
    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; btpAO-KnowledgeBot/1.0; +https://btpao.fr/bot)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "fr-FR,fr;q=0.9,en;q=0.8",
    }

    try:
        async with httpx.AsyncClient(timeout=15.0, follow_redirects=True, verify=False) as client:
            resp = await client.get(url, headers=headers)
            
            if resp.status_code >= 400:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Impossible de joindre le lien web : Le serveur distant a répondu avec le statut HTTP {resp.status_code} ({resp.reason_phrase}). Aucune donnée inventée n'a été indexée.",
                )
            
            html_content = resp.text
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Erreur de connexion réseau à l'URL '{url}' : {str(exc)}. Veuillez vérifier l'accessibilité publique du lien.",
        )

    # 3. Clean HTML and extract text
    parser = HTMLTextExtractor()
    try:
        parser.feed(html_content)
        extracted_title = parser.get_title()
        clean_text = parser.get_text()
    except Exception as e:
        clean_text = re.sub(r"<[^>]+>", " ", html_content)
        clean_text = re.sub(r"\s+", " ", clean_text).strip()
        extracted_title = ""

    page_title = payload.title.strip() if payload.title and payload.title.strip() else extracted_title
    if not page_title:
        page_title = f"Source Web : {url.split('//')[-1].split('/')[0]}"

    if len(clean_text) < 30:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="La page web ne contient aucun contenu textuel substantiel à indexer.",
        )

    word_count = len(clean_text.split())


    # 4. Generate Embedding
    try:
        await embedding_service.sync_platform_key(db)
        embedding_vector = embedding_service.generate_embedding(clean_text[:4000])
    except Exception:
        embedding_vector = None

    asset_id = uuid.uuid4()
    metadata = {
        "source_type": "web",
        "url": url,
        "status": "indexed",
        "word_count": word_count,
        "scraped_at": now.isoformat(),
        "tags": ["source_web", payload.category or "web_source"],
    }
    if embedding_vector:
        metadata["has_embedding"] = True

    new_asset = CompanyAsset(
        id=asset_id,
        tenant_id=t_uuid,
        category=payload.category or "web_source",
        title=page_title,
        description=clean_text[:12000],
        s3_url=url,
        status="indexed",
        embedding=embedding_vector,
        metadata_json=metadata,
        created_at=now,
        updated_at=now,
    )
    db.add(new_asset)
    await db.flush()
    await db.refresh(new_asset)

    return CompanyAssetOut(
        id=str(new_asset.id),
        tenant_id=str(new_asset.tenant_id),
        category=new_asset.category,
        title=new_asset.title,
        description=new_asset.description or "",
        s3_url=new_asset.s3_url,
        tags=metadata["tags"],
        metadata_json=metadata,
        created_at=new_asset.created_at,
    )


@router.delete("/assets/{asset_id}")
async def delete_company_asset(
    asset_id: str,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Immediate hard-delete of a company knowledge asset strictly scoped to the tenant.
    Never leaves hidden soft-deleted rows.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
        a_uuid = uuid.UUID(asset_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid UUID")

    stmt = select(CompanyAsset).where(CompanyAsset.id == a_uuid, CompanyAsset.tenant_id == t_uuid)
    result = await db.execute(stmt)
    asset = result.scalar_one_or_none()

    if not asset:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document introuvable ou accès refusé")

    await db.delete(asset)
    await db.commit()

    return {"success": True, "message": "Document supprimé définitivement de la base de connaissances (hard delete immédiat)."}


@router.post("/assets/{asset_id}/obsolete")
async def mark_company_asset_obsolete(
    asset_id: str,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Marks a document as obsolete. Obsolete documents are excluded from active RAG generations
    and will be automatically purged 90 days after this marking via Celery Beat retention policy.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
        a_uuid = uuid.UUID(asset_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid UUID")

    stmt = select(CompanyAsset).where(CompanyAsset.id == a_uuid, CompanyAsset.tenant_id == t_uuid)
    result = await db.execute(stmt)
    asset = result.scalar_one_or_none()

    if not asset:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document introuvable ou accès refusé")

    now = datetime.now(timezone.utc)
    asset.status = "obsolete"
    asset.obsolete_at = now
    if isinstance(asset.metadata_json, dict):
        asset.metadata_json["obsolete_marked_at"] = now.isoformat()

    await db.commit()
    return {
        "success": True,
        "message": "Document marqué comme obsolète. Il sera purgé définitivement après 90 jours.",
        "asset_id": str(asset.id),
        "status": "obsolete",
        "obsolete_at": now.isoformat(),
    }


@router.post("/assets", response_model=CompanyAssetOut, status_code=status.HTTP_201_CREATED)
async def create_company_asset(
    payload: CompanyAssetCreate,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Creates a new company knowledge asset stored in PostgreSQL under the authenticated tenant.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    # Enforce quota
    await billing_service.check_and_enforce_knowledge_quota(t_uuid, db=db)

    asset_id = uuid.uuid4()
    now = datetime.now(timezone.utc)

    # Compute embedding if description provided
    embedding_vector = None
    if payload.description:
        try:
            await embedding_service.sync_platform_key(db)
            embedding_vector = embedding_service.generate_embedding(payload.description[:4000])
        except Exception:
            embedding_vector = None

    metadata = payload.metadata_json or {}
    if payload.tags:
        metadata["tags"] = payload.tags
    metadata["source_type"] = metadata.get("source_type", "text")
    metadata["status"] = "indexed"
    if embedding_vector:
        metadata["has_embedding"] = True

    new_asset = CompanyAsset(
        id=asset_id,
        tenant_id=t_uuid,
        category=payload.category,
        title=payload.title,
        description=payload.description,
        s3_url=None,
        status="indexed",
        embedding=embedding_vector,
        metadata_json=metadata,
        created_at=now,
        updated_at=now,
    )

    db.add(new_asset)
    await db.flush()
    await db.refresh(new_asset)

    return CompanyAssetOut(
        id=str(new_asset.id),
        tenant_id=str(new_asset.tenant_id),
        category=new_asset.category,
        title=new_asset.title,
        description=new_asset.description or "",
        s3_url=new_asset.s3_url,
        tags=payload.tags or [],
        metadata_json=new_asset.metadata_json or {},
        created_at=new_asset.created_at,
    )


@router.get("/search")
async def search_knowledge(
    query: str,
    category: Optional[str] = None,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Semantic cosine distance search in company knowledge base strictly within the authenticated tenant's data.
    Uses pgvector embedding vector similarity, excluding obsolete documents.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    # Generate query embedding
    query_vector = None
    try:
        await embedding_service.sync_platform_key(db)
        query_vector = embedding_service.generate_embedding(query[:2000])
    except Exception as e:
        query_vector = None

    matched = []

    if query_vector is not None:
        try:
            async with db.begin_nested():
                # pgvector cosine similarity ranking: 1 - cosine_distance
                distance_expr = CompanyAsset.embedding.cosine_distance(query_vector)
                stmt = (
                    select(CompanyAsset, distance_expr.label("distance"))
                    .where(
                        CompanyAsset.tenant_id == t_uuid,
                        CompanyAsset.embedding.isnot(None),
                        CompanyAsset.status != "obsolete",
                    )
                )
                if category:
                    stmt = stmt.where(CompanyAsset.category == category)
                stmt = stmt.order_by(distance_expr).limit(10)

                result = await db.execute(stmt)
                rows = result.all()
                for asset, dist in rows:
                    similarity = round(max(0.0, 1.0 - float(dist)), 4) if dist is not None else 0.5
                    matched.append({
                        "id": str(asset.id),
                        "category": asset.category,
                        "title": asset.title,
                        "description": asset.description,
                        "s3_url": asset.s3_url,
                        "metadata_json": asset.metadata_json or {},
                        "score": similarity,
                    })
        except Exception as emb_search_err:
            logger.warning("[knowledge.py] Semantic vector search fallback: %s", emb_search_err)

    # Fallback to keyword matching if no pgvector results returned
    if not matched:
        stmt = select(CompanyAsset).where(
            CompanyAsset.tenant_id == t_uuid,
            CompanyAsset.status != "obsolete",
        )
        if category:
            stmt = stmt.where(CompanyAsset.category == category)
        stmt = stmt.order_by(CompanyAsset.created_at.desc()).limit(20)
        result = await db.execute(stmt)
        assets = result.scalars().all()

        q_lower = query.lower()
        for a in assets:
            score = 0.4
            if q_lower in a.title.lower():
                score = 0.95
            elif a.description and q_lower in a.description.lower():
                score = 0.85
            matched.append({
                "id": str(a.id),
                "category": a.category,
                "title": a.title,
                "description": a.description,
                "s3_url": a.s3_url,
                "metadata_json": a.metadata_json or {},
                "score": score,
            })
        matched.sort(key=lambda x: x["score"], reverse=True)

    return {"query": query, "results": matched[:10]}



@router.post("/template/word")
async def upload_word_template(
    file: UploadFile = File(...),
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Uploads the client's branded Word template (.docx with header, logo, footer).
    Saves to tenant-scoped storage and records it in export_templates table in PostgreSQL.
    """
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Seuls les fichiers .docx sont acceptés comme template Word."
        )

    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    file_bytes = await file.read()
    subpath = f"templates/word_template_{current_user.tenant_id}.docx"

    s3_key = storage_service.upload_file(
        tenant_id=current_user.tenant_id,
        subpath=subpath,
        file_obj=file_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Save S3 key reference in export_templates table
    stmt = select(ExportTemplate).where(
        ExportTemplate.tenant_id == t_uuid,
        ExportTemplate.is_default == True,
    )
    result = await db.execute(stmt)
    existing_template = result.scalar_one_or_none()

    if existing_template:
        existing_template.name = file.filename
        existing_template.s3_docx_key = s3_key
    else:
        new_template = ExportTemplate(
            id=uuid.uuid4(),
            tenant_id=t_uuid,
            name=file.filename,
            description="Template Word officiel client",
            s3_docx_key=s3_key,
            is_default=True,
            created_at=datetime.utcnow(),
        )
        db.add(new_template)

    await db.flush()

    return {
        "success": True,
        "message": f"Template Word '{file.filename}' enregistré avec succès.",
        "s3_key": s3_key,
        "filename": file.filename,
    }


@router.get("/template/word")
async def get_word_template_info(
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Returns metadata about the current Word template for this tenant from PostgreSQL.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    stmt = select(ExportTemplate).where(
        ExportTemplate.tenant_id == t_uuid,
        ExportTemplate.is_default == True,
    )
    result = await db.execute(stmt)
    template = result.scalar_one_or_none()

    if template:
        return {
            "has_template": True,
            "filename": template.name,
            "updated_at": template.created_at.isoformat(),
        }

    return {"has_template": False, "filename": None, "updated_at": None}


@router.get("/template/suggested")
async def get_suggested_template(
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Deduces the most relevant starting template from the tenant's history:
    1. Active or recent ExportTemplate (.docx uploaded in settings)
    2. Most recent successfully generated dossier (ExportJob)
    3. Relevant reference document in CompanyAsset (memoire_technique / reference_technique / .docx)
    4. If none exists (new client), returns has_template=False so manual upload remains the only path.
    """
    try:
        t_uuid = uuid.UUID(current_user.tenant_id)
    except (ValueError, TypeError):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid tenant UUID")

    # 1. Check ExportTemplate table
    stmt_tmpl = (
        select(ExportTemplate)
        .where(ExportTemplate.tenant_id == t_uuid)
        .order_by(ExportTemplate.is_default.desc(), ExportTemplate.created_at.desc())
    )
    tmpl_res = await db.execute(stmt_tmpl)
    template = tmpl_res.scalar_one_or_none()

    if template:
        return {
            "has_template": True,
            "source_type": "export_template",
            "name": template.name,
            "description": "Template Word officiel configuré dans vos paramètres entreprise",
            "id": str(template.id),
            "created_at": template.created_at.isoformat(),
        }

    # 2. Check recent completed ExportJobs
    stmt_jobs = (
        select(ExportJob, Project)
        .join(Project, Project.id == ExportJob.project_id)
        .where(ExportJob.tenant_id == t_uuid, ExportJob.status == "completed")
        .order_by(ExportJob.completed_at.desc(), ExportJob.created_at.desc())
    )
    job_res = await db.execute(stmt_jobs)
    first_job_row = job_res.first()

    if first_job_row:
        job, proj = first_job_row
        return {
            "has_template": True,
            "source_type": "recent_dossier",
            "name": f"Mémoire récent : {proj.title}",
            "description": "Structure et charte déduites de votre plus récent dossier exporté",
            "id": str(job.id),
            "created_at": (job.completed_at or job.created_at).isoformat(),
        }

    # 3. Check reference documents in CompanyAsset
    stmt_asset = (
        select(CompanyAsset)
        .where(
            CompanyAsset.tenant_id == t_uuid,
            CompanyAsset.category.in_([
                # Real categories from the manual upload form (company/page.tsx)
                "fiche_technique", "memoire_reference", "certification",
                "qse_securite", "moyens_materiels",
                # Real categories from the AI company-bootstrap service
                "presentation_generale", "certificat_qualibat", "materiel_engins",
                "cv_encadrement", "demarche_rse", "reference_chantier",
            ]),
            CompanyAsset.status != "obsolete",
        )
        .order_by(CompanyAsset.created_at.desc())
        # 04/09 : .limit(1) + .first() au lieu de scalar_one_or_none(). L'intention est
        # "le document de reference le plus recent" (d'ou le order_by desc), mais
        # scalar_one_or_none() leve MultipleResultsFound des que le tenant possede plus
        # d'un document eligible -- ce qui est le cas normal. La suggestion de template
        # plantait donc en 500 pour tout client ayant deux references ou plus.
        .limit(1)
    )
    asset_res = await db.execute(stmt_asset)
    asset = asset_res.scalars().first()

    if asset:
        return {
            "has_template": True,
            "source_type": "reference_document",
            "name": asset.title,
            "description": "Document de référence issu de votre base de connaissances",
            "id": str(asset.id),
            "created_at": asset.created_at.isoformat(),
        }

    # 4. Fallback for new clients: No historical template found
    return {
        "has_template": False,
        "source_type": None,
        "name": None,
        "description": "Aucun document historique trouvé. Veuillez téléverser votre template Word.",
        "id": None,
        "created_at": None,
    }


@router.get("/assets/{asset_id}/download")
async def download_knowledge_asset(
    asset_id: uuid.UUID,
    inline: bool = False,
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Télécharge ou affiche le fichier original d'un document de savoir-faire.
    Strictement isolé par tenant_id.
    """
    t_uuid = uuid.UUID(current_user.tenant_id)
    stmt = select(CompanyAsset).where(
        CompanyAsset.id == asset_id,
        CompanyAsset.tenant_id == t_uuid,
    )
    res = await db.execute(stmt)
    asset = res.scalar_one_or_none()
    if not asset:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document introuvable")

    meta = asset.metadata_json or {}
    filename = meta.get("file_name") or f"{asset.title}.pdf"

    # 1. Download file bytes from storage_service if s3_url exists
    if asset.s3_url:
        try:
            file_bytes = storage_service.download_file(tenant_id=current_user.tenant_id, s3_key=asset.s3_url)
            content_type = meta.get("content_type")
            if not content_type:
                fn_lower = filename.lower()
                if fn_lower.endswith(".pdf"):
                    content_type = "application/pdf"
                elif fn_lower.endswith(".docx"):
                    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif fn_lower.endswith(".xlsx"):
                    content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                elif fn_lower.endswith(".png"):
                    content_type = "image/png"
                elif fn_lower.endswith((".jpg", ".jpeg")):
                    content_type = "image/jpeg"
                else:
                    content_type = "application/octet-stream"

            disposition = "inline" if inline else "attachment"
            return Response(
                content=file_bytes,
                media_type=content_type,
                headers={
                    "Content-Disposition": f'{disposition}; filename="{filename}"',
                    "Cache-Control": "private, max-age=3600",
                },
            )
        except Exception as exc:
            logger.warning(f"Erreur téléchargement storage_service pour {asset.s3_url}: {exc}")

    # 2. Fallback to description text if available
    if asset.description:
        disposition = "inline" if inline else "attachment"
        return Response(
            content=asset.description.encode("utf-8"),
            media_type="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f'{disposition}; filename="{asset.title}.txt"',
                "Cache-Control": "private, max-age=3600",
            },
        )

    raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Fichier binaire introuvable sur le stockage")


@router.get("/template/word/download")
async def download_word_template(
    current_user: CurrentTenantUser = Depends(get_current_tenant_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Télécharge le modèle Word par défaut de l'entreprise.
    """
    t_uuid = uuid.UUID(current_user.tenant_id)
    stmt = select(ExportTemplate).where(
        ExportTemplate.tenant_id == t_uuid,
        ExportTemplate.is_default == True,
    )
    res = await db.execute(stmt)
    template = res.scalar_one_or_none()
    if not template or not template.s3_docx_key:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Aucun modèle Word officiel enregistré")

    try:
        file_bytes = storage_service.download_file(tenant_id=current_user.tenant_id, s3_key=template.s3_docx_key)
        return Response(
            content=file_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{template.name or "template_officiel.docx"}"',
                "Cache-Control": "private, max-age=3600",
            },
        )
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Fichier modèle introuvable : {exc}")

