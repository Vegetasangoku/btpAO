"""
MEA Regional Tender Dossiers & RTL OpenXML Word Generator Engine.
Supports Saudi Arabia (SA), Qatar (QA), UAE (AE), and Lebanon (LB) in English and Arabic.
Injects native OpenXML RTL attributes (w:bidi, w:rtl) for Arabic document rendering.
"""
import io
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


class MEADossierService:
    @staticmethod
    def _apply_rtl_to_paragraph(paragraph):
        """
        Injects OpenXML w:bidi and right alignment for RTL text (Arabic).
        """
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    @staticmethod
    def _apply_rtl_to_run(run, font_name: str = "Traditional Arabic"):
        """
        Injects OpenXML w:rtl attribute and Arabic-compatible font.
        """
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        run.font.name = font_name

    def generate_saudi_tender_dossier(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        language: str = "en",
    ) -> bytes:
        """
        Saudi Arabia (SA) - Government Tender & Procurement Law (GTPL) Dossier.
        Includes: Form of Tender, ZATCA/Zakat declaration, GOSI social security,
        MOMRAH Contractor Classification & Nitaqat Saudization statement.
        """
        doc = docx.Document()
        is_ar = language == "ar"
        is_fr = language == "fr"

        if is_ar:
            # Arabic Title & Header
            p_title = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_title)
            r_title = p_title.add_run("المملكة العربية السعودية — ملف العطاء والمنافسات الحكومية")
            self._apply_rtl_to_run(r_title)
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(16, 85, 50)

            p_sub = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_sub)
            r_sub = p_sub.add_run("نظام المنافسات والمشتريات الحكومية (مرسوم ملكي رقم م/128) — منصة اعتماد")
            self._apply_rtl_to_run(r_sub)
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            # Sections in Arabic
            p1 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p1)
            r1 = p1.add_run(f"١. بيانات المقاول: {tenant.get('name', 'شركة المقاولات')} (سجل تجاري: {tenant.get('siret', tenant.get('cr_number', '1010000000'))})")
            self._apply_rtl_to_run(r1)

            p2 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p2)
            r2 = p2.add_run(f"٢. المشروع والجهة المالكة: {project.get('title', 'مشروع مقاولات')} — {project.get('client_name', 'الجهة الحكومية')} (مرجع: {project.get('reference_code', 'SA-2026')})")
            self._apply_rtl_to_run(r2)

            p3 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p3)
            r3 = p3.add_run("٣. الشهادات الإلزامية المسجلة: شهادة الزكاة والدخل (ZATCA)، شهادة التأمينات الاجتماعية (GOSI)، تصنيف المقاولين (بلدي/MOMRAH)، ونسبة التوطين (نطاقات).")
            self._apply_rtl_to_run(r3)

            p4 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p4)
            r4 = p4.add_run("٤. إقرار الالتزام: نقر بالتزامنا بكود البناء السعودي (SBC) ومتطلبات المحتوى المحلي (LCGPA) بنسبة مطابقة كاملة.")
            self._apply_rtl_to_run(r4)
        elif is_fr:
            # Version Française (02/09, correctif tâche #66 : "Français" produisait
            # silencieusement de l'anglais faute de branche dédiée ici)
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("ROYAUME D'ARABIE SAOUDITE — FORMULAIRE OFFICIEL DE SOUMISSION")
            r_title.font.name = "Arial"
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(16, 85, 50)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Loi sur les marchés publics et la concurrence (GTPL — Décret royal n° M/128) — Conformité portail Etimad")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            sa_project_title = project.get('title', 'Travaux de Construction')
            sa_client_name = project.get('client_name', 'Entité Gouvernementale')
            doc.add_paragraph().add_run(f"1. Identification du soumissionnaire : {tenant.get('name', 'Société de Travaux')} (Registre de commerce (CR) : {tenant.get('siret', tenant.get('cr_number', '1010000000'))})").bold = True
            doc.add_paragraph().add_run(f"2. Autorité adjudicatrice et projet : {sa_project_title} — {sa_client_name} (Réf. : {project.get('reference_code', 'SA-2026')})")
            doc.add_paragraph().add_run("3. Attestations légales obligatoires : Certificat de Zakat (ZATCA), assurance sociale (GOSI), classification des entrepreneurs (MOMRAH) et taux de saoudisation (Nitaqat) vérifiés.")
            doc.add_paragraph().add_run("4. Engagement technique : Conformité totale au Code de la construction saoudien (SBC 201/801) et aux normes de contenu local (LCGPA).")
        else:
            # English Version
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("KINGDOM OF SAUDI ARABIA — OFFICIAL FORM OF TENDER")
            r_title.font.name = "Arial"
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(16, 85, 50)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Government Tender & Procurement Law (GTPL - Royal Decree M/128) — Etimad Portal Compliance")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            doc.add_paragraph().add_run(f"1. Contractor Identification: {tenant.get('name', 'Contractor Ltd')} (CR: {tenant.get('siret', tenant.get('cr_number', '1010000000'))})").bold = True
            doc.add_paragraph().add_run(f"2. Tendering Authority & Project: {project.get('title', 'Civil Works')} — {project.get('client_name', 'Government Entity')} (Ref: {project.get('reference_code', 'SA-2026')})")
            doc.add_paragraph().add_run("3. Mandatory Statutory Compliances: ZATCA Zakat Certificate, GOSI Social Insurance, MOMRAH Contractor Classification & Nitaqat Saudization status verified.")
            doc.add_paragraph().add_run("4. Technical Commitment: Full compliance with the Saudi Building Code (SBC 201/801) and Local Content and Government Procurement Authority (LCGPA) standards.")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_qatar_tender_dossier(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        language: str = "en",
    ) -> bytes:
        """
        Qatar (QA) - Tender Law No. 24 of 2015 & Ashghal Directives.
        Includes: Form of Tender, ICV (In-Country Value) commitment, QCS 2018 specifications.
        """
        doc = docx.Document()
        is_ar = language == "ar"
        is_fr = language == "fr"

        if is_ar:
            p_title = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_title)
            r_title = p_title.add_run("دولة قطر — ملف المناقصات والعطاءات الرسمية (أشغال / مناقصات)")
            self._apply_rtl_to_run(r_title)
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(140, 29, 64)

            p_sub = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_sub)
            r_sub = p_sub.add_run("قانون تنظيم المناقصات والمزايدات رقم (٢٤) لسنة ٢٠١٥ والمواصفات العامة للإنشاء (QCS 2018)")
            self._apply_rtl_to_run(r_sub)

            p1 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p1)
            r1 = p1.add_run(f"١. اسم الشركة والمقاول: {tenant.get('name', 'شركة المقاولات')} (السجل التجاري: {tenant.get('siret', 'CR-QAT-001')})")
            self._apply_rtl_to_run(r1)

            p2 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p2)
            r2 = p2.add_run(f"٢. تفاصيل المناقصة: {project.get('title', 'مشروع قطر')} — الجهة: {project.get('client_name', 'هيئة الأشغال العامة أشغال')} (مرجع: {project.get('reference_code', 'QA-2026')})")
            self._apply_rtl_to_run(r2)

            p3 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p3)
            r3 = p3.add_run("٣. معايير القيمة المحلية (ICV): التزام تام بخطة القيمة المحلية المضافة وبرنامج توطين التابع لقطر للطاقة / وزارة المالية.")
            self._apply_rtl_to_run(r3)
        elif is_fr:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("ÉTAT DU QATAR — FORMULAIRE DE SOUMISSION ASHGHAL / MONAQASAT")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(140, 29, 64)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Loi n° 24 de 2015 sur les appels d'offres et Spécifications de construction du Qatar (QCS 2018)")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            qa_project_title = project.get('title', "Travaux d'Infrastructure")
            qa_client_name = project.get('client_name', "Autorité des Travaux Publics (Ashghal)")
            doc.add_paragraph().add_run(f"1. Coordonnées du soumissionnaire : {tenant.get('name', 'Société de Travaux')} (N° de registre de commerce : {tenant.get('siret', 'CR-QAT-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Marché et maître d'ouvrage : {qa_project_title} — {qa_client_name} (Réf. : {project.get('reference_code', 'QA-2026')})")
            doc.add_paragraph().add_run("3. Valeur en pays (ICV) : Score de valeur ajoutée locale certifié et documentation d'achat local jointe.")
            doc.add_paragraph().add_run("4. Conformité aux normes : Exécution stricte selon le QCS 2018 et la certification de durabilité GSAS 4 étoiles.")
        else:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("STATE OF QATAR — ASHGHAL / MONAQASAT FORM OF TENDER")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(140, 29, 64)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Tender Law No. 24 of 2015 & Qatar Construction Specifications (QCS 2018)")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            doc.add_paragraph().add_run(f"1. Tenderer Details: {tenant.get('name', 'Contractor Ltd')} (CR No: {tenant.get('siret', 'CR-QAT-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Project & Client: {project.get('title', 'Infrastructure Works')} — {project.get('client_name', 'Public Works Authority (Ashghal)')} (Ref: {project.get('reference_code', 'QA-2026')})")
            doc.add_paragraph().add_run("3. In-Country Value (ICV): Certified In-Country Value score and local procurement compliance documentation attached.")
            doc.add_paragraph().add_run("4. Standards Compliance: Strict execution according to QCS 2018 and GSAS 4-Star Sustainability Rating.")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_uae_tender_dossier(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        language: str = "en",
    ) -> bytes:
        """
        United Arab Emirates (UAE) - Federal Procurement Law No. 11/2023 & MoF / DED Directives.
        """
        doc = docx.Document()
        is_ar = language == "ar"
        is_fr = language == "fr"

        if is_ar:
            p_title = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_title)
            r_title = p_title.add_run("دولة الإمارات العربية المتحدة — نموذج العطاء الاتحادي")
            self._apply_rtl_to_run(r_title)
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(0, 115, 47)

            p_sub = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_sub)
            r_sub = p_sub.add_run("المرسوم بقانون اتحادي رقم (١١) لسنة ٢٠٢٣ في شأن المشتريات الحكومية — بوابة المشتريات الاتحادية")
            self._apply_rtl_to_run(r_sub)

            p1 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p1)
            r1 = p1.add_run(f"١. المنشأة المتقدمة: {tenant.get('name', 'شركة المقاولات')} (الرخصة التجارية DED: {tenant.get('siret', 'CN-UAE-001')})")
            self._apply_rtl_to_run(r1)

            p2 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p2)
            r2 = p2.add_run(f"٢. المناقصة: {project.get('title', 'مشروع الإنشاءات')} — {project.get('client_name', 'الجهة الاتحادية / المحلية')} (مرجع: {project.get('reference_code', 'UAE-2026')})")
            self._apply_rtl_to_run(r2)

            p3 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p3)
            r3 = p3.add_run("٣. نسبة التوطين ومعايير الاستدامة: الالتزام بقرارات وزارة الموارد البشرية والتوطين (MOHRE) ومعايير استدامة (Estidama Pearl / Al Sa'fat).")
            self._apply_rtl_to_run(r3)
        elif is_fr:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("ÉMIRATS ARABES UNIS — FORMULAIRE DE SOUMISSION FÉDÉRAL")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(0, 115, 47)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Décret-loi fédéral n° 11 de 2023 relatif aux marchés publics — Plateforme numérique du Ministère des Finances")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            uae_project_title = project.get('title', 'Projet de Développement')
            uae_client_name = project.get('client_name', "Ministère de l'Énergie et des Infrastructures")
            doc.add_paragraph().add_run(f"1. Identification de l'entreprise : {tenant.get('name', 'Société de Travaux LLC')} (Licence commerciale : {tenant.get('siret', 'CN-UAE-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Entité adjudicatrice et projet : {uae_project_title} — {uae_client_name} (Réf. : {project.get('reference_code', 'UAE-2026')})")
            doc.add_paragraph().add_run("3. Émiratisation et conformité sociale : Conformité totale aux quotas d'émiratisation du MOHRE et au système de protection des salaires (WPS).")
            doc.add_paragraph().add_run("4. Normes de construction durable : Conformité certifiée à la notation Estidama Pearl (Abou Dabi) et à la réglementation Al Sa'fat (Dubaï).")
        else:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("UNITED ARAB EMIRATES — FEDERAL FORM OF TENDER")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(0, 115, 47)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Federal Decree-Law No. 11 of 2023 on Federal Government Procurement — MoF Digital Platform")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            doc.add_paragraph().add_run(f"1. Tenderer Details: {tenant.get('name', 'Contractor LLC')} (Trade License: {tenant.get('siret', 'CN-UAE-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Procurement Entity & Project: {project.get('title', 'Development Project')} — {project.get('client_name', 'Ministry of Energy & Infrastructure')} (Ref: {project.get('reference_code', 'UAE-2026')})")
            doc.add_paragraph().add_run("3. Emiratisation & Labour Compliance: Full compliance with MOHRE Emiratisation quotas and Wage Protection System (WPS).")
            doc.add_paragraph().add_run("4. Green Building Codes: Certified compliance with Estidama Pearl Rating (Abu Dhabi) & Dubai Green Building Regulations (Al Sa'fat).")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_lebanon_tender_dossier(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        language: str = "fr",
    ) -> bytes:
        """
        Lebanon (LB) - Public Procurement Authority (PPA - Loi n° 244/2021).
        Supports French and Arabic.
        """
        doc = docx.Document()
        is_ar = language == "ar"
        is_en = language == "en"

        if is_ar:
            p_title = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_title)
            r_title = p_title.add_run("الجمهورية اللبنانية — هيئة الشراء العام (PPA)")
            self._apply_rtl_to_run(r_title)
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(237, 27, 36)

            p_sub = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p_sub)
            r_sub = p_sub.add_run("ملف تقديم العروض وفقاً لأحكام قانون الشراء العام رقم ٢٤٤/٢٠٢١")
            self._apply_rtl_to_run(r_sub)

            p1 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p1)
            r1 = p1.add_run(f"١. العارض: {tenant.get('name', 'شركة المقاولات')} (السجل التجاري: {tenant.get('siret', 'RC-BEY-001')})")
            self._apply_rtl_to_run(r1)

            p2 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p2)
            r2 = p2.add_run(f"٢. الصفقة: {project.get('title', 'أشغال')} — الإدارة المتعاقدة: {project.get('client_name', 'مجلس الإنماء والإعمار / وزارة الأشغال')} (مرجع: {project.get('reference_code', 'LB-2026')})")
            self._apply_rtl_to_run(r2)

            p3 = doc.add_paragraph()
            self._apply_rtl_to_paragraph(p3)
            r3 = p3.add_run("٣. المستندات الإدارية: براءة ذمة من الصندوق الوطني للضمان الاجتماعي (CNSS)، إفادة تسجيل لدى وزارة المالية، وتأشيرة نقابة المهندسين (OIA).")
            self._apply_rtl_to_run(r3)
        elif is_en:
            # 02/09, correctif tâche #66 : bug miroir du cas SA/QA/AE -- ici c'était
            # "English" qui produisait silencieusement du français (le else ci-dessous
            # n'a jamais été que la version française, jamais anglaise).
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("LEBANESE REPUBLIC — PUBLIC PROCUREMENT AUTHORITY (PPA)")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(237, 27, 36)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Bid submission dossier in accordance with Law No. 244/2021 on public procurement")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            doc.add_paragraph().add_run(f"1. Bidder: {tenant.get('name', 'Contracting Company SAL')} (Commercial Registry: {tenant.get('siret', 'RC-BEY-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Contracting Authority & Project: {project.get('title', 'Public Works')} — {project.get('client_name', 'Council for Development and Reconstruction (CDR)')} (Ref.: {project.get('reference_code', 'LB-2026')})")
            doc.add_paragraph().add_run("3. Legal Certificates: Tax clearance from the Ministry of Finance, NSSF certificate, and visas from the Order of Engineers and Architects of Beirut/Tripoli (OEA).")
            doc.add_paragraph().add_run("4. Sworn Declaration: No bankruptcy, no conflict of interest, and full legal compliance under Articles 14 to 18 of Law 244/2021.")
        else:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = p_title.add_run("RÉPUBLIQUE LIBANAISE — AUTORITÉ DES MARCHÉS PUBLICS (PPA)")
            r_title.font.size = Pt(16)
            r_title.font.bold = True
            r_title.font.color.rgb = RGBColor(237, 27, 36)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_sub = p_sub.add_run("Dossier de candidature et soumission conforme à la Loi n° 244/2021 sur les marchés publics")
            r_sub.font.size = Pt(10)
            r_sub.font.italic = True

            doc.add_paragraph().add_run(f"1. Candidat : {tenant.get('name', 'Entreprise SAL')} (RCS : {tenant.get('siret', 'RC-BEY-001')})").bold = True
            doc.add_paragraph().add_run(f"2. Autorité adjudicatrice & Projet : {project.get('title', 'Travaux Publics')} — {project.get('client_name', 'Conseil du Développement et de la Reconstruction (CDR)')} (Réf : {project.get('reference_code', 'LB-2026')})")
            doc.add_paragraph().add_run("3. Attestations Légales : Quitus fiscal Ministère des Finances, bordeaux CNSS et visas de l'Ordre des Ingénieurs et Architectes de Beyrouth/Tripoli (OIA).")
            doc.add_paragraph().add_run("4. Déclaration sur l'honneur : Non-faillite, absence de conflit d'intérêts et régularité juridique complète selon les articles 14 à 18 de la Loi 244/2021.")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


mea_dossier_service = MEADossierService()
