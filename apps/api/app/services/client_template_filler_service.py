"""
Client Template In-Place Word (.docx) Filler & Completeness Checklist Service.
Preserves existing Word styles, fonts, margins, headers and tables without alteration.
Strict Multi-Tier Source Hierarchy & Anti-Hallucination [À COMPLÉTER] red flags.
"""
import io
import logging
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple
import docx
from docx.shared import Pt, RGBColor
from pydantic import BaseModel

logger = logging.getLogger("client_template_filler")

RED_ALERT_COLOR = RGBColor(220, 38, 38) # Tailwind red-600


class CompletenessItem(BaseModel):
    section_name: str
    status: str # 'filled', 'action_required'
    source_used: Optional[str] = None
    missing_elements: List[str] = []


class CompletenessReport(BaseModel):
    total_fields: int
    filled_fields: int
    pending_actions_count: int
    completeness_score_pct: float
    is_ready_for_submission: bool
    sections: List[CompletenessItem]
    generated_at: str


class ClientTemplateFillerService:
    @staticmethod
    def _resolve_field_value(
        field_key: str,
        project_data: Dict[str, Any],
        rag_chunks: List[str],
        company_assets: Dict[str, Any],
        tenant_learnings: List[Dict[str, Any]],
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Multi-tier Source Hierarchy Resolution:
        1. Project / DCE Explicit Data
        2. RAG Semantic Search Knowledge
        3. Validated Company Assets
        4. Tenant Historical Learnings
        Returns (value, source_tier_name) or (None, None)
        """
        key_lower = field_key.lower().strip()

        # Tier 1: Project / DCE Data
        if key_lower in ("client_name", "acheteur", "maitre_ouvrage") and project_data.get("client_name"):
            return str(project_data["client_name"]), "Tier 1: DCE / Projet Explicite"
        if key_lower in ("project_title", "titre_marche", "operation") and project_data.get("title"):
            return str(project_data["title"]), "Tier 1: DCE / Projet Explicite"
        if key_lower in ("reference_code", "reference", "consultation") and project_data.get("reference_code"):
            return str(project_data["reference_code"]), "Tier 1: DCE / Projet Explicite"
        if key_lower in ("lot_number", "lot") and project_data.get("lot_number"):
            return str(project_data["lot_number"]), "Tier 1: DCE / Projet Explicite"

        # Tier 2: RAG Semantic Search
        for chunk in rag_chunks:
            if key_lower in chunk.lower() and len(chunk.strip()) > 10:
                return chunk.strip()[:400], "Tier 2: RAG Sémantique DCE"

        # Tier 3: Validated Company Assets
        if key_lower in ("siret", "siren") and company_assets.get("siret"):
            return str(company_assets["siret"]), "Tier 3: Asset Entreprise Validé"
        if key_lower in ("company_name", "raison_sociale", "entreprise") and company_assets.get("name"):
            return str(company_assets["name"]), "Tier 3: Asset Entreprise Validé"
        if key_lower in ("headcount", "effectif") and company_assets.get("headcount"):
            return str(company_assets["headcount"]), "Tier 3: Asset Entreprise Validé"
        if key_lower in ("insurance", "assurance", "decennale") and company_assets.get("insurance"):
            return str(company_assets["insurance"]), "Tier 3: Asset Entreprise Validé"

        # Tier 4: Tenant Historical Learnings
        for learning in tenant_learnings:
            directive = learning.get("directive", "") or learning.get("insight", "")
            if key_lower in directive.lower():
                return directive, "Tier 4: Apprentissage Entreprise Validé"

        return None, None

    def fill_docx_template_inplace(
        self,
        template_bytes: bytes,
        project_data: Dict[str, Any],
        rag_chunks: Optional[List[str]] = None,
        company_assets: Optional[Dict[str, Any]] = None,
        tenant_learnings: Optional[List[Dict[str, Any]]] = None,
    ) -> Tuple[bytes, CompletenessReport]:
        """
        Fills placeholders inside an uploaded client DOCX template without altering styles.
        Recognizes {{placeholder}}, [CHAMP], <<field>> patterns.
        Unresolved fields receive bold red [À COMPLÉTER : ...] flags.
        """
        doc = docx.Document(io.BytesIO(template_bytes))
        rag = rag_chunks or []
        assets = company_assets or {}
        learnings = tenant_learnings or []

        pattern = re.compile(r"\{\{([^\}]+)\}\}|<<([^>]+)>>|\[\[([^\]]+)\]\]")

        total_fields = 0
        filled_fields = 0
        completeness_items: List[CompletenessItem] = []

        def process_paragraph(paragraph):
            nonlocal total_fields, filled_fields
            full_text = paragraph.text
            matches = list(pattern.finditer(full_text))
            if not matches:
                return

            for match in matches:
                raw_key = match.group(1) or match.group(2) or match.group(3)
                key = raw_key.strip()
                total_fields += 1

                resolved_val, source_tier = self._resolve_field_value(key, project_data, rag, assets, learnings)

                if resolved_val:
                    filled_fields += 1
                    # Replace in text while preserving paragraph formatting
                    for run in paragraph.runs:
                        if match.group(0) in run.text:
                            run.text = run.text.replace(match.group(0), resolved_val)
                    completeness_items.append(
                        CompletenessItem(
                            section_name=key,
                            status="filled",
                            source_used=source_tier,
                            missing_elements=[],
                        )
                    )
                else:
                    # Insert red [À COMPLÉTER] flag
                    missing_desc = f"[À COMPLÉTER : {key}]"
                    for run in paragraph.runs:
                        if match.group(0) in run.text:
                            run.text = run.text.replace(match.group(0), missing_desc)
                            run.font.color.rgb = RED_ALERT_COLOR
                            run.font.bold = True
                    completeness_items.append(
                        CompletenessItem(
                            section_name=key,
                            status="action_required",
                            source_used=None,
                            missing_elements=[f"Donnée manquante : {key}"],
                        )
                    )

        # Process standard body paragraphs
        for p in doc.paragraphs:
            process_paragraph(p)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        process_paragraph(p)

        # Calculate metrics
        pending = total_fields - filled_fields
        score = round((filled_fields / total_fields * 100), 1) if total_fields > 0 else 100.0
        is_ready = pending == 0

        report = CompletenessReport(
            total_fields=total_fields,
            filled_fields=filled_fields,
            pending_actions_count=pending,
            completeness_score_pct=score,
            is_ready_for_submission=is_ready,
            sections=completeness_items,
            generated_at=datetime.now(timezone.utc).isoformat(),
        )

        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        return out_buffer.getvalue(), report


client_template_filler_service = ClientTemplateFillerService()
