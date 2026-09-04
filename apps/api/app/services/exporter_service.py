"""
Word (.docx) and PDF Document Exporter Engine
Assemble le document via python-docx (pas de moteur de templating Jinja2/docxtpl --
l'import ci-dessous a été retiré le 02/09 car réellement inutilisé, cf. investigation
tâche #66), insertion directe des images Gantt & Organigramme, et LibreOffice CLI pour
la conversion PDF headless (réellement appelée depuis le 02/09, voir convert_docx_to_pdf
et son unique appelant dans app/workers/tasks.py::build_export_doc_task).
"""
import io
import os
import re
import subprocess
import tempfile
import unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional
import docx
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from app.core.storage import storage_service
from app.services.gantt_service import gantt_service
from app.services.diagram_service import diagram_service


class ExporterService:
    # Libelles fixes ("chrome") du document Word genere -- traduits une fois ici plutot
    # que d'inserer un if/elif par langue a chaque endroit du document (30/08). Les
    # sections elles-memes (content_html) sont deja generees dans la bonne langue par
    # llm_generator.py ; ceci ne couvre que la mise en forme fixe autour (page de garde,
    # sommaire, legendes de figures, avertissements).
    _EXPORT_I18N = {
        "fr": {
            "default_company": "Votre Entreprise",
            "title_badge": "RÉPONSE À L'APPEL D'OFFRES",
            "main_title": "MÉMOIRE TECHNIQUE JUSTIFICATIF",
            "default_project_title": "Projet BTP",
            "ref_label": "Réf. Consultation :",
            "moa_label": "Maître d'Ouvrage :",
            "lot_label": "Lot :",
            "default_ref": "AO-2026",
            "default_client": "Client",
            "default_lot": "Lot 01",
            "row_delai": "Délai global d'exécution garanti :",
            "unit_mois": "mois",
            "row_budget": "Budget prévisionnel des travaux :",
            "row_materiel": "Matériel lourd principal :",
            "default_materiel": "Grue Potain 50m",
            "row_dechets": "Taux de valorisation déchets BTP :",
            "val_dechets": "88% en filières locales agréées (<15 km)",
            "toc_heading": "Sommaire du Mémoire Technique",
            "missing_sections_warning": "⚠️ Sections requises par l'appel d'offres, absentes de la structure du template client d'origine (ajoutées automatiquement ci-dessous, à vérifier) :",
            "figure1_caption": "\nFigure 1 : Organigramme d'Encadrement Chantier",
            "figure2_caption": "\nFigure 2 : Planning Prévisionnel de Phasage (Gantt)",
            "chart_warning_prefix": "⚠️ Avertissement : Génération du graphique échouée — ",
            "gantt_error_msg": "Le planning prévisionnel (Gantt) n'a pas pu être généré automatiquement pour cet export.",
            "organigramme_error_msg": "L'organigramme d'encadrement n'a pas pu être généré automatiquement pour cet export.",
            "default_section_title": "Section",
        },
        "en": {
            "default_company": "Your Company",
            "title_badge": "RESPONSE TO THE TENDER",
            "main_title": "TECHNICAL PROPOSAL",
            "default_project_title": "Construction Project",
            "ref_label": "Tender Reference:",
            "moa_label": "Contracting Authority:",
            "lot_label": "Lot:",
            "default_ref": "RFP-2026",
            "default_client": "Client",
            "default_lot": "Lot 01",
            "row_delai": "Guaranteed overall completion time:",
            "unit_mois": "months",
            "row_budget": "Estimated works budget:",
            "row_materiel": "Main heavy equipment:",
            "default_materiel": "Potain 50m tower crane",
            "row_dechets": "Construction waste recovery rate:",
            "val_dechets": "88% through approved local channels (<15 km)",
            "toc_heading": "Table of Contents",
            "missing_sections_warning": "⚠️ Sections required by the tender, absent from the original client template structure (automatically added below, to be reviewed):",
            "figure1_caption": "\nFigure 1: Site Supervision Organization Chart",
            "figure2_caption": "\nFigure 2: Preliminary Phasing Schedule (Gantt)",
            "chart_warning_prefix": "⚠️ Warning: Chart generation failed — ",
            "gantt_error_msg": "The preliminary schedule (Gantt) could not be generated automatically for this export.",
            "organigramme_error_msg": "The supervision organization chart could not be generated automatically for this export.",
            "default_section_title": "Section",
        },
        "ar": {
            "default_company": "شركتكم",
            "title_badge": "الرد على طلب العروض",
            "main_title": "المذكرة الفنية التبريرية",
            "default_project_title": "مشروع بناء",
            "ref_label": "مرجع الاستشارة:",
            "moa_label": "صاحب المشروع:",
            "lot_label": "الحصة:",
            "default_ref": "AO-2026",
            "default_client": "العميل",
            "default_lot": "الحصة 01",
            "row_delai": "المدة الإجمالية المضمونة للتنفيذ:",
            "unit_mois": "أشهر",
            "row_budget": "الميزانية التقديرية للأشغال:",
            "row_materiel": "المعدات الثقيلة الرئيسية:",
            "default_materiel": "رافعة برجية Potain بارتفاع 50 م",
            "row_dechets": "معدل تثمين نفايات البناء:",
            "val_dechets": "88% عبر قنوات محلية معتمدة (أقل من 15 كم)",
            "toc_heading": "فهرس المذكرة الفنية",
            "missing_sections_warning": "⚠️ أقسام مطلوبة في طلب العروض وغير موجودة في نموذج العميل الأصلي (أُضيفت تلقائيًا أدناه، يُرجى المراجعة):",
            "figure1_caption": "\nالشكل 1: الهيكل التنظيمي لتأطير الورش",
            "figure2_caption": "\nالشكل 2: المخطط الزمني التقديري للمراحل (Gantt)",
            "chart_warning_prefix": "⚠️ تنبيه: فشل توليد الرسم البياني — ",
            "gantt_error_msg": "تعذر توليد المخطط الزمني التقديري (Gantt) تلقائيًا لهذا التصدير.",
            "organigramme_error_msg": "تعذر توليد الهيكل التنظيمي للتأطير تلقائيًا لهذا التصدير.",
            "default_section_title": "قسم",
        },
    }

    def build_memo_docx(
        self,
        tenant_id: str,
        project_id: str,
        project_data: Dict[str, Any],
        sections: List[Dict[str, Any]],
        decision_form: Dict[str, Any],
        template_bytes: Optional[bytes] = None,
        include_visuals: bool = True,
        include_cover_page: bool = True,
        required_section_titles: Optional[List[str]] = None,
        gantt_tasks: Optional[List[Dict[str, Any]]] = None,
        language: str = "fr",
        brand_color: Optional[str] = None,
        shape_style: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Compiles all validated sections, metadata, Gantt planning, and organigramme into a Word .docx document.
        Uses the client's Word template (with their header/logo/footer) when provided.
        `brand_color` / `shape_style` (branding_config.primary_color / .shape_style, BT02
        01/09) are forwarded to the Gantt & organigramme generators so the exported Word
        visuals match the client's own color charter and shape preset instead of always
        rendering with the hardcoded defaults (previously the case for every export --
        callers never passed brand_color here even though visuals.py already had it for
        the web preview).
        """
        EXP = self._EXPORT_I18N.get(language, self._EXPORT_I18N["fr"])

        # 1. Create a base document — use client template if available
        missing_sections: List[str] = []
        if template_bytes:
            try:
                doc = docx.Document(io.BytesIO(template_bytes))
                missing_sections = self._detect_missing_required_sections(template_bytes, required_section_titles)
                # Clear existing body content but preserve headers/footers
                for element in list(doc.element.body):
                    if element.tag.endswith('}sectPr'):
                        continue  # preserve section properties (margins, headers)
                    doc.element.body.remove(element)
                self._replace_company_placeholders(doc, project_data.get('company_name') or EXP['default_company'])
            except Exception as e:
                print(f"[ExporterService] Template load error, using blank: {e}")
                doc = docx.Document()
        else:
            doc = docx.Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # 2. Cover Page Header (correctif tâche #66, 02/09 : le toggle "Page de
        # garde" existait déjà côté UI (page.tsx) mais n'avait jamais été câblé
        # jusqu'ici -- ce bloc s'exécutait toujours inconditionnellement.
        if include_cover_page:
            p_header = doc.add_paragraph()
            r_logo = p_header.add_run(f"{project_data.get('company_name') or EXP['default_company']}\n")
            r_logo.bold = True
            r_logo.font.size = Pt(16)
            r_logo.font.color.rgb = RGBColor(2, 132, 199)

            # Title
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("\n" * 2)
        
            r_title_badge = p_title.add_run(f"{EXP['title_badge']}\n")
            r_title_badge.font.size = Pt(13)
            r_title_badge.font.color.rgb = RGBColor(100, 116, 139)
            r_title_badge.bold = True

            r_main_title = p_title.add_run(f"{EXP['main_title']}\n{project_data.get('title', EXP['default_project_title']).upper()}\n")
            r_main_title.font.size = Pt(22)
            r_main_title.bold = True
            r_main_title.font.color.rgb = RGBColor(15, 23, 42)

            p_ref = doc.add_paragraph()
            p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_ref = p_ref.add_run(f"{EXP['ref_label']} {project_data.get('reference_code', EXP['default_ref'])}\n{EXP['moa_label']} {project_data.get('client_name', EXP['default_client'])}\n{EXP['lot_label']} {project_data.get('lot_number', EXP['default_lot'])}\n")
            r_ref.font.size = Pt(12)
            r_ref.font.color.rgb = RGBColor(51, 65, 85)

            doc.add_paragraph("\n" * 2)

            # Executive Summary Box Table
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            summary_table.style = 'Light Shading Accent 1' if 'Light Shading Accent 1' in [s.name for s in doc.styles] else 'Table Grid'

            rows_data = [
                (EXP['row_delai'], f"{decision_form.get('delai_mois', 6)} {EXP['unit_mois']}"),
                (EXP['row_budget'], f"{project_data.get('budget_estimate', 3500000.0):,.2f} € HT"),
                (EXP['row_materiel'], decision_form.get('materiel_principal', EXP['default_materiel'])),
                (EXP['row_dechets'], EXP['val_dechets']),
            ]

            for idx, (label, val) in enumerate(rows_data):
                row = summary_table.rows[idx]
                c1 = row.cells[0].paragraphs[0].add_run(label)
                c1.bold = True
                row.cells[1].paragraphs[0].add_run(val)

            doc.add_page_break()

        # 3. Table of Contents & Sections
        p_toc = doc.add_heading(EXP['toc_heading'], level=1)
        p_toc.paragraph_format.space_after = Pt(14)

        for s in sections:
            p_toc_item = doc.add_paragraph(style='List Bullet')
            r_toc = p_toc_item.add_run(f"{s.get('title', 'Section')}")
            r_toc.font.size = Pt(11)

        if missing_sections:
            warn_p = doc.add_paragraph()
            warn_run = warn_p.add_run(EXP['missing_sections_warning'])
            warn_run.bold = True
            warn_run.font.size = Pt(9.5)
            warn_run.font.color.rgb = RGBColor(180, 83, 9)
            for missing_title in missing_sections:
                li = doc.add_paragraph(style='List Bullet')
                li_run = li.add_run(missing_title)
                li_run.font.size = Pt(9.5)
                li_run.font.color.rgb = RGBColor(180, 83, 9)

        doc.add_paragraph("\n")

        # 4. Generate Visuals if requested
        gantt_path = None
        organigramme_path = None
        gantt_error = None
        organigramme_error = None

        if include_visuals:
            try:
                if gantt_tasks:
                    # Real, possibly user-edited tasks from the interactive Gantt (Batch 11)
                    # take priority over the legacy static phase list, so what the user sees
                    # and edits on screen is what actually gets embedded in the Word export.
                    gantt_res = gantt_service.generate_gantt_chart_png_from_tasks(
                        tenant_id=tenant_id,
                        project_id=project_id,
                        project_title=project_data.get("title", "Chantier BTP"),
                        tasks=gantt_tasks,
                        brand_color=brand_color,
                        shape_style=shape_style,
                    )
                else:
                    gantt_res = gantt_service.generate_gantt_chart_png(
                        tenant_id=tenant_id,
                        project_id=project_id,
                        project_title=project_data.get("title", "Chantier BTP"),
                        phases=decision_form.get("phasage_travaux", []),
                        brand_color=brand_color,
                        shape_style=shape_style,
                    )
                gantt_bytes = storage_service.download_file(tenant_id, gantt_res["s3_key"])
                temp_gantt = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                temp_gantt.write(gantt_bytes)
                temp_gantt.close()
                gantt_path = temp_gantt.name
            except Exception as e:
                print(f"[ExporterService] Gantt generation error: {e}")
                gantt_error = EXP['gantt_error_msg']

            try:
                diag_res = diagram_service.generate_organigramme_png(
                    tenant_id=tenant_id,
                    project_id=project_id,
                    project_title=project_data.get("title", "Chantier BTP"),
                    cadres=decision_form.get("equipe_cadres", []),
                    brand_color=brand_color,
                    shape_style=shape_style,
                )
                diag_bytes = storage_service.download_file(tenant_id, diag_res["s3_key"])
                temp_diag = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                temp_diag.write(diag_bytes)
                temp_diag.close()
                organigramme_path = temp_diag.name
            except Exception as e:
                print(f"[ExporterService] Organigramme generation error: {e}")
                organigramme_error = EXP['organigramme_error_msg']

        # 5. Render Each Section Body
        for s in sections:
            sec_key = s.get("section_key", "")
            title = s.get("title", EXP['default_section_title'])
            html_content = s.get("content_html", "")

            h = doc.add_heading(title, level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(10)

            # Convert HTML text to clean paragraphs
            clean_paragraphs = self._html_to_docx_paragraphs(html_content)
            for p_text in clean_paragraphs:
                p_elem = doc.add_paragraph()
                p_elem.paragraph_format.line_spacing = 1.15
                p_elem.paragraph_format.space_after = Pt(6)
                
                # Check for bullet points
                if p_text.startswith("• ") or p_text.startswith("- "):
                    p_elem.style = 'List Bullet'
                    p_text = p_text[2:]
                
                run = p_elem.add_run(p_text)
                run.font.size = Pt(10.5)

            # Insert Visuals at appropriate sections
            if sec_key == "moyens_humains":
                if organigramme_path and os.path.exists(organigramme_path):
                    doc.add_paragraph(EXP['figure1_caption']).runs[0].italic = True
                    doc.add_picture(organigramme_path, width=Inches(6.5))
                    doc.add_paragraph("\n")
                elif organigramme_error:
                    warn_p = doc.add_paragraph()
                    warn_run = warn_p.add_run(f"{EXP['chart_warning_prefix']}{organigramme_error}")
                    warn_run.italic = True
                    warn_run.font.color.rgb = RGBColor(185, 28, 28)

            elif sec_key == "methodologie_phasage" or sec_key == "planning_gantt":
                if gantt_path and os.path.exists(gantt_path):
                    doc.add_paragraph(EXP['figure2_caption']).runs[0].italic = True
                    doc.add_picture(gantt_path, width=Inches(6.5))
                    doc.add_paragraph("\n")
                elif gantt_error:
                    warn_p = doc.add_paragraph()
                    warn_run = warn_p.add_run(f"{EXP['chart_warning_prefix']}{gantt_error}")
                    warn_run.italic = True
                    warn_run.font.color.rgb = RGBColor(185, 28, 28)

        # 5bis. Direction RTL pour l'arabe -- OOXML n'expose aucune propriete RTL de haut
        # niveau via python-docx ; passage final unique sur tout le document plutot que
        # de modifier chaque add_paragraph/add_heading individuellement ci-dessus (30/08)
        if language == "ar":
            self._apply_rtl_to_document(doc)

        # 6. Save docx to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_bytes = docx_buffer.read()

        # Clean up temporary visual files
        for p in [gantt_path, organigramme_path]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass

        # 7. Attempt to save to tenant storage (non-blocking)
        s3_docx_key = f"tenants/{tenant_id}/exports/{project_id}/memoire_technique_{project_id[:8]}.docx"
        try:
            storage_service.upload_file(
                tenant_id=tenant_id,
                subpath=f"exports/{project_id}/memoire_technique_{project_id[:8]}.docx",
                file_obj=docx_bytes,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            print(f"[ExporterService] Storage upload notice (file still available in-memory): {e}")

        return {
            "s3_docx_key": s3_docx_key,
            "docx_bytes": docx_bytes,
            "filename": f"Memoire_Technique_{project_data.get('reference_code', 'AO')}.docx",
        }

    def convert_docx_to_pdf(self, docx_bytes: bytes, tenant_id: str, project_id: str) -> Optional[str]:
        """
        Converts DOCX bytes to PDF using LibreOffice headless.
        Falls back gracefully if LibreOffice is not installed in the environment.
        """
        with tempfile.TemporaryDirectory() as tmp_dir:
            in_file = os.path.join(tmp_dir, "input.docx")
            with open(in_file, "wb") as f:
                f.write(docx_bytes)

            try:
                cmd = ["libreoffice", "--headless", "--convert-to", "pdf", in_file, "--outdir", tmp_dir]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30)
                
                out_pdf = os.path.join(tmp_dir, "input.pdf")
                if os.path.exists(out_pdf):
                    with open(out_pdf, "rb") as f:
                        pdf_bytes = f.read()

                    s3_pdf_key = storage_service.upload_file(
                        tenant_id=tenant_id,
                        subpath=f"exports/{project_id}/memoire_technique_{project_id[:8]}.pdf",
                        file_obj=pdf_bytes,
                        content_type="application/pdf"
                    )
                    return s3_pdf_key
            except Exception as e:
                print(f"[ExporterService] LibreOffice PDF conversion notice: {e}")

        return None

    # ─── Template Structural Fidelity (Batch 9) ────────────────────────────
    # AO-required sections missing from an uploaded client template are
    # detected (not silently dropped) and flagged as a visible suggestion; a
    # company-name placeholder left in the template's own header/footer is
    # replaced with the real tenant name instead of shipping whatever was
    # hardcoded in the client's original file.

    _PLACEHOLDER_PATTERNS = [
        re.compile(r'\{\{?\s*nom[_ ]du[_ ]client\s*\}?\}', re.IGNORECASE),
        re.compile(r'\{\{?\s*nom[_ ]entreprise\s*\}?\}', re.IGNORECASE),
        re.compile(r'\[\s*NOM[_ ]DU[_ ]CLIENT\s*\]', re.IGNORECASE),
        re.compile(r'\[\s*NOM[_ ]ENTREPRISE\s*\]', re.IGNORECASE),
    ]

    def _replace_placeholder_in_paragraphs(self, paragraphs, company_name: str) -> int:
        replaced = 0
        for p in paragraphs:
            for run in p.runs:
                new_text = run.text
                for pattern in self._PLACEHOLDER_PATTERNS:
                    new_text, n = pattern.subn(company_name, new_text)
                    replaced += n
                if new_text != run.text:
                    run.text = new_text
        return replaced

    def _replace_company_placeholders(self, doc, company_name: str) -> int:
        """
        Best-effort: replaces {nom_du_client}-style placeholder tokens found in the
        template's headers/footers with the real tenant company name. The generated
        body always replaces the template's own body content wholesale (see
        build_memo_docx below), so only headers/footers can still carry a hardcoded
        name from the client's original file -- this is the one place that needs it.
        Does not scan first-page/even-page header-footer variants (uncommon in
        practice); a token split across multiple Word-internal runs is not caught
        (documented limitation, non-blocking).
        """
        total = 0
        for section in doc.sections:
            for hf in (section.header, section.footer):
                try:
                    total += self._replace_placeholder_in_paragraphs(hf.paragraphs, company_name)
                    for table in hf.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                total += self._replace_placeholder_in_paragraphs(cell.paragraphs, company_name)
                except Exception as e:
                    print(f"[ExporterService] Placeholder replacement notice: {e}")
        return total

    @staticmethod
    def _normalize_heading_text(text: str) -> str:
        text = unicodedata.normalize('NFKD', text)
        text = ''.join(c for c in text if not unicodedata.combining(c))
        text = re.sub(r'^[0-9]+[\.\)]?\s*', '', text)
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        return ' '.join(text.lower().split())

    @staticmethod
    def _word_set_similarity(norm_a: str, norm_b: str) -> float:
        """
        Jaccard similarity over normalized word sets. Deliberately NOT a raw
        character-level ratio (e.g. difflib.SequenceMatcher): short French BTP
        section titles share common domain words -- "Moyens Humains &
        Encadrement" vs "Moyens Matériels & Engins" overlap heavily at the
        character level (shared "Moyens", shared letters) despite being
        completely different sections. Comparing whole-word sets instead means
        only "Moyens" overlaps (1 of 5 unique words, ratio 0.2), correctly
        telling them apart, while still tolerating minor real-world phrasing
        differences (a template heading missing "& Encadrement" still matches
        at 0.67).
        """
        words_a, words_b = set(norm_a.split()), set(norm_b.split())
        if not words_a or not words_b:
            return 0.0
        return len(words_a & words_b) / len(words_a | words_b)

    def _detect_missing_required_sections(
        self, template_bytes: Optional[bytes], required_titles: Optional[List[str]]
    ) -> List[str]:
        """
        Compares the heading-styled paragraph text found in the uploaded client
        template against the AO-required canonical section titles. Returns the
        canonical titles with no reasonably similar heading in the template, so
        build_memo_docx can flag them as a visible suggestion instead of silently
        inserting content the template's own author never anticipated. Returns []
        when there is no template (nothing to compare against) or no required
        titles were supplied -- never invents a "missing" list out of nothing.
        """
        if not template_bytes or not required_titles:
            return []
        try:
            tpl_doc = docx.Document(io.BytesIO(template_bytes))
        except Exception as e:
            print(f"[ExporterService] Template heading scan notice: {e}")
            return []

        heading_texts = []
        for p in tpl_doc.paragraphs:
            text = p.text.strip()
            if not text or len(text) > 120:
                continue
            style_name = p.style.name if p.style else ''
            is_heading_style = 'Heading' in style_name or 'Title' in style_name
            is_bold_run = any(r.bold for r in p.runs if r.bold)
            if is_heading_style or is_bold_run:
                heading_texts.append(self._normalize_heading_text(text))

        missing = []
        for title in required_titles:
            norm_title = self._normalize_heading_text(title)
            best_ratio = max(
                (self._word_set_similarity(norm_title, h) for h in heading_texts),
                default=0.0,
            )
            if best_ratio < 0.5:
                missing.append(title)
        return missing

    def _apply_rtl_to_document(self, doc) -> None:
        """
        Applique la direction RTL (droite-a-gauche) a l'ensemble du document Word --
        necessaire pour un rendu correct en arabe (ordre de lecture, alignement,
        numerotation de listes, ordre des colonnes de tableau). python-docx n'expose
        aucune propriete RTL de haut niveau ; manipulation directe de l'OOXML sous-jacent
        (w:bidi / w:bidiVisual), seule approche fiable avec cette bibliotheque.

        Applique en un seul passage final sur tout le document (corps + cellules de
        tableau + sections) plutot qu'a chaque add_paragraph/add_heading individuel dans
        build_memo_docx -- plus sur (un seul endroit a faire fonctionner correctement) et
        couvre aussi les paragraphes deja presents dans un template client fourni. (30/08)
        """
        def _rtl_paragraph(p) -> None:
            p_pr = p._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            p_pr.append(bidi)
            if p.alignment in (None, WD_ALIGN_PARAGRAPH.LEFT):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for p in doc.paragraphs:
            _rtl_paragraph(p)

        for table in doc.tables:
            try:
                tbl_pr = table._tbl.tblPr
                tbl_bidi = OxmlElement('w:bidiVisual')
                tbl_pr.append(tbl_bidi)
            except Exception as e:
                print(f"[ExporterService] RTL table notice: {e}")
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _rtl_paragraph(p)

        for section in doc.sections:
            try:
                sect_pr = section._sectPr
                sect_bidi = OxmlElement('w:bidi')
                sect_pr.append(sect_bidi)
            except Exception as e:
                print(f"[ExporterService] RTL section notice: {e}")

    def _html_to_docx_paragraphs(self, html_text: str) -> List[str]:
        """
        Strips HTML tags while preserving paragraph breaks and bullet list structures.
        """
        text = html_text.replace("</li>", "\n")
        text = text.replace("<li>", "• ")
        text = text.replace("</h2>", "\n\n")
        text = text.replace("</h3>", "\n\n")
        text = text.replace("</p>", "\n\n")
        text = text.replace("<br>", "\n")
        text = text.replace("<br/>", "\n")
        
        # Remove remaining tags
        clean = re.sub(r'<[^>]+>', '', text)
        
        # Clean entities
        clean = clean.replace("&nbsp;", " ").replace("&amp;", "&").replace("&quot;", '"').replace("&#39;", "'").replace("&gt;", ">").replace("&lt;", "<")
        
        paragraphs = [p.strip() for p in clean.split("\n\n") if p.strip()]
        return paragraphs


exporter_service = ExporterService()
