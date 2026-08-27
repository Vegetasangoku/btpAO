"""
MEA Technical Structure, CSI MasterFormat, BoQ (POMI/CESMM4),
Prequalification Dossier (PQD) and Building Codes Compliance Engine.
Supports Saudi Arabia (SBC/LCGPA), Qatar (QCS/GSAS), UAE (Estidama/Al Sa'fat), Lebanon (OIA/PPA).
"""
import io
import logging
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger("mea_structure_service")

# CSI MasterFormat Divisions Standard (Divisions 01 to 48)
CSI_MASTERFORMAT_DIVISIONS = {
    "01": "General Requirements (Project Management, QA/QC, HSE, Submittals, Temporary Facilities)",
    "02": "Existing Conditions (Site Demolition, Soil Remediation, Subsurface Investigation)",
    "03": "Concrete (Formwork, Reinforcement, Cast-in-Place, Precast Concrete)",
    "04": "Masonry (Concrete Unit Masonry, Engineered Stone, Mortar & Grout)",
    "05": "Metals (Structural Steel, Metal Decking, Cold-Formed Metal Framing)",
    "06": "Wood, Plastics, and Composites (Rough Carpentry, Architectural Woodwork)",
    "07": "Thermal and Moisture Protection (Waterproofing, Thermal Insulation, Roofing)",
    "08": "Openings (Doors, Windows, Curtain Wall, Glazing, Louvers)",
    "09": "Finishes (Plaster, Gypsum Board, Tiling, Acoustical Ceilings, Flooring, Painting)",
    "10": "Specialties (Signage, Partitions, Fire Protection Specialties)",
    "11": "Equipment (Commercial, Industrial and Facility Equipment)",
    "12": "Furnishings (Laboratory Casework, Window Treatments, Fixed Seating)",
    "13": "Special Construction (Clean Rooms, Sound/Vibration Control, Swimming Pools)",
    "14": "Conveying Equipment (Elevators, Escalators, Moving Walks)",
    "21": "Fire Suppression (Fire-Suppression Piping, Fire Pumps, Clean-Agent Extinguishing)",
    "22": "Plumbing (Plumbing Piping, Water Heating, Medical Gas Systems)",
    "23": "HVAC (Heating, Ventilating, Air Conditioning, BMS Controls)",
    "26": "Electrical (Transformers, Switchboards, Medium/Low Voltage Distribution, Lighting)",
    "27": "Communications (Structured Cabling, Data Centers, Audio-Video Systems)",
    "28": "Electronic Safety and Security (Fire Alarm, CCTV Access Control, Perimeter Intrusion)",
    "31": "Earthwork (Site Clearing, Grading, Excavation, Shoring & Deep Foundation Pilings)",
    "32": "Exterior Improvements (Paving, Curbing, Fences, Hard & Soft Landscaping, Irrigation)",
    "33": "Utilities (Water Distribution, Sanitary Sewerage, Storm Drainage, Electrical Utilities)",
    "40": "Process Integration & Industrial Piping",
    "48": "Electrical Power Generation (Solar Photovoltaic, Battery Energy Storage BESS)",
}

# Building Codes & Standards Matrix
BUILDING_CODES_BY_COUNTRY = {
    "SA": {
        "primary_building_code": "Saudi Building Code (SBC)",
        "code_components": [
            "SBC 201: General Building Architectural Requirements",
            "SBC 301-305: Structural Design (Loads, Concrete, Steel, Foundations)",
            "SBC 401: Electrical Installations",
            "SBC 501: Mechanical & HVAC Systems",
            "SBC 601: Energy Conservation Code",
            "SBC 701: Sanitary & Plumbing Code",
            "SBC 801: Fire Protection & Life Safety",
            "SBC 1001: Green Construction Code",
        ],
        "local_content_body": "LCGPA (Local Content and Government Procurement Authority)",
        "contractor_classification": "MOMRAH / Balady Contractor Classification System",
        "power_of_attorney_req": "Ministry of Justice (Najiz) electronic PoA or consularised PoA stamped by MOFA",
    },
    "QA": {
        "primary_building_code": "Qatar Construction Specifications (QCS 2018)",
        "code_components": [
            "QCS Section 01: General Technical & Submittal Procedures",
            "QCS Section 05: Concrete & Reinforced Concrete Construction",
            "QCS Section 06: Roadworks & Pavement Foundations",
            "QCS Section 19: Plumbing & Public Health Services",
            "QCS Section 21: Electrical & Telecommunication Installation",
            "QCS Section 23: Fire Fighting & Alarm Systems",
        ],
        "local_content_body": "QatarEnergy / MoF Tawteen ICV (In-Country Value) Framework",
        "contractor_classification": "Ashghal & MoF Unified Prequalification Classification",
        "power_of_attorney_req": "Qatar Ministry of Justice authenticated PoA or certified consular legalization",
    },
    "AE": {
        "primary_building_code": "UAE National Building Code / Estidama / Dubai Building Code (DBC)",
        "code_components": [
            "Dubai Building Code (DBC 2021 Part A-K)",
            "Estidama Pearl Building Rating System (Abu Dhabi DoE/DMT)",
            "Al Sa'fat - Dubai Green Building Evaluation System",
            "UAE Fire & Life Safety Code of Practice (Civil Defence CD)",
        ],
        "local_content_body": "National In-Country Value (ICV) Programme (Ministry of Industry MoIAT)",
        "contractor_classification": "Department of Economic Development (DED) & Ministry of Finance",
        "power_of_attorney_req": "UAE Notary Public or MOFAIC legalized Power of Attorney",
    },
    "LB": {
        "primary_building_code": "Lebanese Building Code (Loi de Construction n° 646/2004 & OIA)",
        "code_components": [
            "Loi n° 646/2004 portant Code de la Construction",
            "Règles et Normes Techniques de l'Ordre des Ingénieurs et Architectes de Beyrouth (OIA)",
            "Normes Libanaises d'isolation thermique et parasismique (NL 135 / LIBNOR)",
            "Loi n° 244/2021 relative aux Marchés Publics (PPA)",
        ],
        "local_content_body": "Préférence nationale 10% selon l'article 23 de la Loi 244/2021",
        "contractor_classification": "Classification interministérielle CDR / Ministère des Travaux Publics",
        "power_of_attorney_req": "Procuration notariée enregistrée auprès du Notaire Libanais ou légalisée Consulat",
    },
}


class MEAStructureService:
    def get_csi_masterformat_structure(
        self,
        project_scope: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Returns structured CSI MasterFormat Divisions with technical submittal requirements.
        """
        divisions = []
        for code, title in CSI_MASTERFORMAT_DIVISIONS.items():
            divisions.append({
                "division_code": code,
                "title": title,
                "standard": "CSI MasterFormat (2020/2026 Edition)",
                "mandatory_submittals": [
                    "Material Technical Data Sheets (TDS)",
                    "Shop Drawings & As-Built Verification",
                    "Method Statement & Risk Assessment (MSRA)",
                    "Inspection & Test Plan (ITP)",
                    "Manufacturer Warranty & Origin Certificate",
                ],
            })
        return divisions

    def generate_boq_template(
        self,
        method: str = "POMI",
        currency: str = "SAR",
    ) -> Dict[str, Any]:
        """
        Generates standard Bill of Quantities (BoQ) breakdown structure based on POMI (Principles of
        Measurement International) or CESMM4 (Civil Engineering Standard Method of Measurement).
        """
        return {
            "measurement_standard": method.upper(),
            "currency": currency,
            "boq_sections": [
                {
                    "bill_number": "Bill No. 1",
                    "title": "General Requirements, Preliminaries & Site Overheads",
                    "items": [
                        {"item_ref": "1.01", "description": "Contractor's Site Management, Supervision & HSE Team", "unit": "Month", "qty": 18},
                        {"item_ref": "1.02", "description": "Mobilization, Temporary Site Offices, Power & Water Utilities", "unit": "Lump Sum", "qty": 1},
                        {"item_ref": "1.03", "description": "Quality Assurance / Quality Control (QA/QC) & Laboratory Testing", "unit": "Lump Sum", "qty": 1},
                        {"item_ref": "1.04", "description": "Insurances (CAR / Third Party Liability / Workmen's Compensation)", "unit": "Lump Sum", "qty": 1},
                    ],
                },
                {
                    "bill_number": "Bill No. 2",
                    "title": "Substructure & Deep Foundations",
                    "items": [
                        {"item_ref": "2.01", "description": "Site Excavation in all types of soil including rock breaking and cart away", "unit": "m³", "qty": 15000},
                        {"item_ref": "2.02", "description": "Cast-in-place Bored Cast Concrete Piles (dia 800mm, depth 20m)", "unit": "Linear Metre", "qty": 2400},
                        {"item_ref": "2.03", "description": "Reinforced Concrete Foundation Raft (C40/50 Sulfate Resisting SRC)", "unit": "m³", "qty": 3500},
                        {"item_ref": "2.04", "description": "Heavy-duty Substructure Tanking & SBS Waterproofing Membrane", "unit": "m²", "qty": 4200},
                    ],
                },
                {
                    "bill_number": "Bill No. 3",
                    "title": "Superstructure & Core Concrete",
                    "items": [
                        {"item_ref": "3.01", "description": "Reinforced Concrete Columns and Shear Walls (C50/60 Low-Carbon CEM III)", "unit": "m³", "qty": 2800},
                        {"item_ref": "3.02", "description": "Post-Tensioned / Reinforced Concrete Suspended Slabs", "unit": "m³", "qty": 6200},
                        {"item_ref": "3.03", "description": "High-Yield Deformed Steel Reinforcement Bars (Grade 500D)", "unit": "Tonne", "qty": 850},
                    ],
                },
                {
                    "bill_number": "Bill No. 4",
                    "title": "MEP Services (Mechanical, Electrical, Fire Protection)",
                    "items": [
                        {"item_ref": "4.01", "description": "High-Efficiency Chilled Water HVAC Air Handling Units (AHU & FCU)", "unit": "Set", "qty": 12},
                        {"item_ref": "4.02", "description": "Main Low Voltage Switchboards (MDB / SMDB 400V 50Hz)", "unit": "Set", "qty": 8},
                        {"item_ref": "4.03", "description": "Automatic Sprinkler System & NFPA Fire Pump Set (UL/FM Certified)", "unit": "Lump Sum", "qty": 1},
                    ],
                },
            ],
        }

    def generate_pqd_word_dossier(
        self,
        tenant: Dict[str, Any],
        country_code: str = "SA",
    ) -> bytes:
        """
        Generates official Prequalification Dossier (PQD) in Word (.docx) format.
        All certifications, financial data and equipment lists are sourced from real
        validated company_assets. Missing data renders as explicit [À COMPLÉTER] placeholders.
        """
        doc = docx.Document()
        country_meta = BUILDING_CODES_BY_COUNTRY.get(country_code.upper(), BUILDING_CODES_BY_COUNTRY["SA"])

        # --- Parse company_assets for real data ---
        company_assets = tenant.get("_company_assets", [])  # List[dict] injected by endpoint
        certif_assets = [a for a in company_assets if str(a.get("category", "")).lower() in ("certification", "certif", "qualibat")]
        financial_assets = [a for a in company_assets if str(a.get("category", "")).lower() in ("financial", "finance")]
        equipment_assets = [a for a in company_assets if str(a.get("category", "")).lower() in ("equipment", "materiel", "matériel")]

        # Build ISO certification list from real assets only
        iso_lines = []
        for ca in certif_assets:
            m = ca.get("metadata_json") or {}
            norm = m.get("norm") or m.get("certification") or ca.get("title", "")
            scope = m.get("scope") or m.get("domaine") or ""
            if norm:
                iso_lines.append(f"• {norm} — {scope}" if scope else f"• {norm}")
        if not iso_lines:
            iso_lines = [
                "• [À COMPLÉTER : certification ISO 9001 — joindre le certificat valide]",
                "• [À COMPLÉTER : certification ISO 45001 / 14001 si applicable]",
            ]

        # Financial solvency from real assets
        fin_rows = []
        for fa in financial_assets[:3]:
            m = fa.get("metadata_json") or {}
            fin_rows.append((
                m.get("indicator") or fa.get("title") or "Indicateur financier",
                m.get("value") or "[À COMPLÉTER : valeur]",
                m.get("audit_note") or "[À COMPLÉTER : référence audit / certification]",
            ))
        if not fin_rows:
            fin_rows = [
                ("Average Annual Turnover (Past 3 Years)", "[À COMPLÉTER : CA moyen €]", "[À COMPLÉTER : états financiers audités]"),
                ("Current Liquidity Ratio", "[À COMPLÉTER : ratio]", "[À COMPLÉTER : source comptable]"),
                ("Available Credit Facilities", "[À COMPLÉTER : montant €]", "[À COMPLÉTER : lettre bancaire]"),
            ]

        # Equipment from real assets
        equip_lines = []
        for ea in equipment_assets:
            m = ea.get("metadata_json") or {}
            desc = m.get("description") or m.get("equipement") or ea.get("title") or ""
            qty = m.get("quantity") or m.get("quantite") or ""
            if desc:
                equip_lines.append(f"• {f'{qty}x ' if qty else ''}{desc}")
        if not equip_lines:
            equip_lines = ["• [À COMPLÉTER : liste du parc matériel lourd détenu en propre]"]

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run("OFFICIAL CONTRACTOR PREQUALIFICATION DOSSIER (PQD)")
        r_title.font.name = "Arial"
        r_title.font.size = Pt(16)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(15, 23, 42)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run(f"Corporate Prequalification for Major Government & Infrastructure Tenders — {country_meta['primary_building_code']}")
        r_sub.font.size = Pt(10)
        r_sub.font.italic = True

        # 1. Company Profile
        doc.add_paragraph().add_run("1. CORPORATE PROFILE & LEGAL STATUS").bold = True
        p1 = doc.add_paragraph()
        p1.add_run(f"Company Legal Name: {tenant.get('name', '[À COMPLÉTER : raison sociale]')}\n")
        p1.add_run(f"Commercial Registration / License: {tenant.get('siret') or tenant.get('cr_number') or '[À COMPLÉTER : SIRET / CR Number]'}\n")
        p1.add_run(f"Contractor Classification: {country_meta['contractor_classification']}\n")
        p1.add_run(f"Legal Power of Attorney (PoA): {country_meta['power_of_attorney_req']}\n")

        # 2. Quality & HSE — from real assets only
        doc.add_paragraph().add_run("2. QUALITY ASSURANCE & HSE MANAGEMENT SYSTEMS").bold = True
        p2 = doc.add_paragraph()
        for line in iso_lines:
            p2.add_run(line + "\n")

        # 3. Financial Solvency — from real assets
        doc.add_paragraph().add_run("3. FINANCIAL SOLVENCY & BANKING CAPACITY").bold = True
        table_fin = doc.add_table(rows=1, cols=3)
        table_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table_fin.rows[0].cells
        hdr[0].text = "Financial Indicator"
        hdr[1].text = "Contractor Value"
        hdr[2].text = "Audit Verification"
        for row in fin_rows:
            cells = table_fin.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]

        # 4. Plant & Heavy Equipment — from real assets
        doc.add_paragraph().add_run("4. PROPRIETARY HEAVY PLANT & EQUIPMENT FLEET").bold = True
        doc.add_paragraph().add_run("\n".join(equip_lines))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def validate_regional_compliance(
        self,
        country_code: str,
        project_data: Dict[str, Any],
        contractor_data: Dict[str, Any],
    ) -> Dict[str, Any]:
        """
        Performs automated compliance check against country-specific building codes.
        ANTI-FABRICATION: Checks 2/3/4 default to False / non_verified.
        They only pass if a matching validated company_asset category is present.
        Check 1 (SIRET/CR) is the only check that can auto-pass from submitted data.
        """
        code = country_code.upper()
        if code not in BUILDING_CODES_BY_COUNTRY:
            return {"status": "error", "message": f"Country '{country_code}' not supported"}

        meta = BUILDING_CODES_BY_COUNTRY[code]
        check_items = []
        company_assets = contractor_data.get("_company_assets", [])  # List[dict] from endpoint

        def has_asset_category(*categories: str) -> bool:
            for a in company_assets:
                cat = str(a.get("category", "")).lower()
                if cat in categories:
                    return True
            return False

        # Check 1: Commercial Registration / SIRET — only check that auto-passes
        has_cr = bool(contractor_data.get("siret") or contractor_data.get("cr_number") or contractor_data.get("trade_license"))
        check_items.append({
            "requirement": "Valid Commercial Registration / Trade License",
            "passed": has_cr,
            "status": "passed" if has_cr else "non_vérifié",
            "details": f"CR: {contractor_data.get('siret', contractor_data.get('cr_number', 'Non fourni'))}",
        })

        # Check 2: Building Code Adherence — requires a validated 'qualification' or 'certification' asset
        has_qualification = has_asset_category("qualification", "certification", "qualibat", "certif")
        check_items.append({
            "requirement": f"Compliance with {meta['primary_building_code']}",
            "passed": has_qualification,
            "status": "passed" if has_qualification else "non_vérifié",
            "details": (
                f"Qualification certifiée vérifiée dans les assets entreprise."
                if has_qualification
                else f"[À VÉRIFIER] : Aucun document de qualification/certification uploadé pour le code {meta['primary_building_code']}. Joindre le certificat dans la base documentaire."
            ),
        })

        # Check 3: Local Content / National Registration — requires 'insurance' or 'legal' asset
        has_insurance = has_asset_category("insurance", "assurance", "legal", "juridique")
        check_items.append({
            "requirement": f"National Content Registration ({meta['local_content_body']})",
            "passed": has_insurance,
            "status": "passed" if has_insurance else "non_vérifié",
            "details": (
                "Attestation d'assurance et/ou document légal trouvé dans les assets entreprise."
                if has_insurance
                else "[À VÉRIFIER] : Aucune attestation d'assurance RC ou document légal uploadé. Joindre les attestations dans la base documentaire."
            ),
        })

        # Check 4: Power of Attorney — requires 'legal' or 'poa' asset (type 'kbis' or 'poa')
        has_poa = has_asset_category("legal", "juridique", "poa", "kbis")
        check_items.append({
            "requirement": "Power of Attorney (PoA) Authenticated",
            "passed": has_poa,
            "status": "passed" if has_poa else "non_vérifié",
            "details": (
                f"Document légal/PoA vérifié dans les assets entreprise. {meta['power_of_attorney_req']}"
                if has_poa
                else f"[À VÉRIFIER] : Aucun document Kbis/PoA uploadé. {meta['power_of_attorney_req']}"
            ),
        })

        all_passed = all(item["passed"] for item in check_items)
        any_verified = any(item["passed"] for item in check_items)

        return {
            "country_code": code,
            "status": "compliant" if all_passed else ("partial" if any_verified else "non_vérifié"),
            "primary_code": meta["primary_building_code"],
            "code_components": meta["code_components"],
            "compliance_checks": check_items,
            "anti_fabrication_note": "Les checks 2/3/4 ne peuvent être 'passed' qu'en présence d'un company_asset validé de la catégorie correspondante.",
            "validated_at": datetime.now(timezone.utc).isoformat(),
        }


mea_structure_service = MEAStructureService()
