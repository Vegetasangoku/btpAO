"""
Service de Génération des Dossiers Administratifs Légaux (DC1, DC2, Récapitulatif DUME).
Permet l'export direct en format Word OpenXML (.docx) et JSON d'interopérabilité Chorus Pro.
"""
import io
import uuid
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class AdminDossierService:
    @staticmethod
    def _create_styled_header(doc: docx.Document, title: str, subtitle: str):
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_title = p_title.add_run(title)
        r_title.font.name = "Arial"
        r_title.font.size = Pt(16)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(14, 116, 144)

        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.name = "Arial"
        r_sub.font.size = Pt(10)
        r_sub.font.italic = True
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()

    @staticmethod
    def _add_section_heading(doc: docx.Document, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)

    def generate_dc1_docx(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        groupement: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        doc = docx.Document()

        self._create_styled_header(
            doc,
            title="FORMULAIRE DC1 — LETTRE DE CANDIDATURE",
            subtitle="Déclaration de candidature et habilitation du mandataire par ses co-traitants (CCP 2026)"
        )

        self._add_section_heading(doc, "A - IDENTIFICATION DE L'ACHETEUR (POUVOIR ADJUDICATEUR)")
        p_a = doc.add_paragraph()
        p_a.add_run("Nom de l'acheteur : ").bold = True
        p_a.add_run(f"{project.get('client_name', 'Non spécifié')}")
        p_a_ref = doc.add_paragraph()
        p_a_ref.add_run("Référence de la consultation : ").bold = True
        p_a_ref.add_run(f"{project.get('reference_code', 'AO-2026')}")

        self._add_section_heading(doc, "B - OBJET DE LA CONSULTATION & DU MARCHÉ")
        p_b = doc.add_paragraph()
        p_b.add_run("Intitulé de l'opération : ").bold = True
        p_b.add_run(f"{project.get('title', 'Marché de travaux BTP')}")
        p_b_lot = doc.add_paragraph()
        p_b_lot.add_run("Allotissement : ").bold = True
        p_b_lot.add_run(f"{project.get('lot_number', 'Lot unique / Offre globale')}")

        self._add_section_heading(doc, "C - FORME DU CANDIDAT (INDIVIDUEL OU GROUPEMENT)")
        p_c = doc.add_paragraph()
        if groupement and groupement.get("is_groupement"):
            p_c.add_run("Type de candidature : Groupement momentané d'entreprises (GME)").bold = True
            p_c_f = doc.add_paragraph()
            p_c_f.add_run(f"Forme : {groupement.get('forme', 'Conjoint avec mandataire solidaire')}")
            p_c_m = doc.add_paragraph()
            p_c_m.add_run(f"Mandataire désigné : {tenant.get('name', 'Notre Entreprise')}")
        else:
            p_c.add_run("Type de candidature : Candidat individuel (Entreprise unique)").bold = True
            p_c_r = doc.add_paragraph()
            p_c_r.add_run(f"Raison sociale : {tenant.get('name', 'Entreprise SAS')}")
            p_c_s = doc.add_paragraph()
            p_c_s.add_run(f"Numéro SIRET : {tenant.get('siret', 'Non renseigné')}")

        self._add_section_heading(doc, "D - DÉCLARATIONS SUR L'HONNEUR DU CANDIDAT")
        p_d = doc.add_paragraph()
        p_d.add_run("Le soussigné déclare sur l'honneur :")
        p_d1 = doc.add_paragraph()
        p_d1.add_run("1. Ne pas faire l'objet d'une interdiction de soumissionner aux marchés publics (Articles L. 2141-1 à L. 2141-11 du Code de la commande publique).")
        p_d2 = doc.add_paragraph()
        p_d2.add_run("2. Être en règle au 31 décembre de l'année précédente au regard de ses obligations fiscales et sociales.")
        p_d3 = doc.add_paragraph()
        p_d3.add_run("3. Que les renseignements fournis dans le présent formulaire et ses annexes sont exacts et sincères.")

        self._add_section_heading(doc, "E - DÉSIGNATION DU SIGNATAIRE HABILITÉ")
        p_e = doc.add_paragraph()
        p_e.add_run(f"Fait à : {tenant.get('city', 'Paris')}, le {datetime.now(timezone.utc).strftime('%d/%m/%Y')}")
        p_e_sig = doc.add_paragraph()
        p_e_sig.add_run("Nom et qualité du signataire : Représentant légal habilité")
        p_e_cachet = doc.add_paragraph()
        p_e_cachet.add_run("Signature & Cachet de l'entreprise : [Signé numériquement]")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_dc2_docx(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
        financial_history: Optional[List[Dict[str, Any]]] = None,
    ) -> bytes:
        doc = docx.Document()

        self._create_styled_header(
            doc,
            title="FORMULAIRE DC2 — DÉCLARATION DU CANDIDAT",
            subtitle="Déclaration des capacités économiques, financières, techniques et professionnelles (CCP 2026)"
        )

        self._add_section_heading(doc, "A - IDENTIFICATION DE L'ENTREPRISE CANDIDATE")
        p_a = doc.add_paragraph()
        p_a.add_run("Dénomination sociale : ").bold = True
        p_a.add_run(f"{tenant.get('name') or '[À COMPLÉTER : dénomination sociale]'}")
        p_a_siret = doc.add_paragraph()
        p_a_siret.add_run("Numéro SIRET : ").bold = True
        p_a_siret.add_run(f"{tenant.get('siret') or '[À COMPLÉTER : numéro SIRET]'}")
        p_a_naf = doc.add_paragraph()
        p_a_naf.add_run("Code NAF / APE : ").bold = True
        p_a_naf.add_run(f"{tenant.get('naf') or '[À COMPLÉTER : code NAF / APE]'}")
        p_a_form = doc.add_paragraph()
        p_a_form.add_run("Forme juridique : ").bold = True
        p_a_form.add_run(f"{tenant.get('legal_form') or '[À COMPLÉTER : forme juridique (ex: SAS, SARL)]'}")

        self._add_section_heading(doc, "B - CAPACITÉS ÉCONOMIQUES ET FINANCIÈRES (CHIFFRE D'AFFAIRES)")
        table_ca = doc.add_table(rows=1, cols=3)
        table_ca.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table_ca.rows[0].cells
        hdr_cells[0].text = "Exercice Comptable"
        hdr_cells[1].text = "Chiffre d'Affaires Global (€ HT)"
        hdr_cells[2].text = "CA Spécifique Marchés Publics / Travaux (€ HT)"

        default_finances = financial_history or [
            {"annee": "[À COMPLÉTER : Exercice N-1]", "ca_global": "[À COMPLÉTER : CA global €]", "ca_specifique": "[À COMPLÉTER : CA marchés publics €]"},
            {"annee": "[À COMPLÉTER : Exercice N-2]", "ca_global": "[À COMPLÉTER : CA global €]", "ca_specifique": "[À COMPLÉTER : CA marchés publics €]"},
            {"annee": "[À COMPLÉTER : Exercice N-3]", "ca_global": "[À COMPLÉTER : CA global €]", "ca_specifique": "[À COMPLÉTER : CA marchés publics €]"},
        ]
        for item in default_finances:
            row_cells = table_ca.add_row().cells
            row_cells[0].text = str(item.get("annee", ""))
            row_cells[1].text = str(item.get("ca_global", ""))
            row_cells[2].text = str(item.get("ca_specifique", ""))

        self._add_section_heading(doc, "C - MOYENS HUMAINS ET TECHNIQUES")
        p_c = doc.add_paragraph()
        p_c.add_run("Effectif moyen annuel permanent : ").bold = True
        p_c.add_run(f"{tenant.get('headcount') or '[À COMPLÉTER : effectif moyen annuel permanent]'}")
        p_c_eq = doc.add_paragraph()
        p_c_eq.add_run("Outillage et matériel lourd détenu en propre : ").bold = True
        p_c_eq.add_run(f"{tenant.get('equipment') or '[À COMPLÉTER : outillage et matériel lourd détenu en propre]'}")

        self._add_section_heading(doc, "D - ATTESTATIONS D'ASSURANCES PROFESSIONNELLES")
        p_d = doc.add_paragraph()
        p_d.add_run("Garantie Décennale & Responsabilité Civile Professionnelle : À justifier par attestation en cours de validité.")
        p_d_cie = doc.add_paragraph()
        insurance_co = tenant.get("insurance_company") or "[À COMPLÉTER : nom de la compagnie d'assurance]"
        p_d_cie.add_run(f"Compagnie d'assurance : {insurance_co}")
        p_d_pol = doc.add_paragraph()
        insurance_pol = tenant.get("insurance_policy_number") or "[À COMPLÉTER : numéro de police d'assurance]"
        p_d_pol.add_run(f"Numéro de police : {insurance_pol}")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_dume_summary(
        self,
        tenant: Dict[str, Any],
        project: Dict[str, Any],
    ) -> Dict[str, Any]:
        """
        Generates structured DUME/ESPD summary for interoperability with Chorus Pro.
        ANTI-FABRICATION: exclusion_grounds and selection_criteria are NEVER auto-declared.
        All boolean fields are None (non-renseigné) — they MUST be filled and validated
        by a legally authorized representative before submission.
        """
        return {
            "dume_version": "ESPD-EDM-V2.1.1",
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "economic_operator": {
                "name": tenant.get("name", "[À COMPLÉTER : dénomination sociale]"),
                "siret": tenant.get("siret", "[À COMPLÉTER : numéro SIRET]"),
                "vat_number": tenant.get("vat_number", f"FR{tenant.get('siret', '000000000')[:9]}" if tenant.get("siret") else "[À COMPLÉTER : numéro TVA intracommunautaire]"),
                "country": tenant.get("country_code", "FR"),
                "is_sme": tenant.get("is_sme", None),  # Must be declared by operator
                "contact": {
                    "email": tenant.get("contact_email", "[À COMPLÉTER : email de contact]"),
                    "phone": tenant.get("contact_phone", "[À COMPLÉTER : téléphone]"),
                },
            },
            "procurement_procedure": {
                "buyer_name": project.get("client_name", "[À COMPLÉTER : nom du pouvoir adjudicateur]"),
                "reference_code": project.get("reference_code", "[À COMPLÉTER : référence de la consultation]"),
                "title": project.get("title", "[À COMPLÉTER : intitulé de l'opération]"),
            },
            # --- DÉCLARATIONS SUR L'HONNEUR : ne jamais pré-remplir ---
            # Chaque champ doit être coché/déclaré par le représentant légal habilité.
            # None = non-déclaré. True = déclaré conforme. False = déclaré non-conforme.
            "exclusion_grounds": {
                "criminal_convictions": None,           # Art. 57(1) Dir. 2014/24/UE — à déclarer manuellement
                "payment_of_taxes": None,               # Art. 57(2) — à déclarer manuellement
                "social_security_contributions": None,  # Art. 57(2) — à déclarer manuellement
                "bankruptcy_insolvency": None,          # Art. 57(4)(b) — à déclarer manuellement
            },
            "selection_criteria": {
                "turnover_requirements_met": None,           # À vérifier sur justificatifs comptables
                "technical_personnel_available": None,       # À justifier par liste nominative
                "references_verified": None,                 # À vérifier sur références signées
                "quality_assurance_schemes": tenant.get("certifications", []),  # From real assets only
            },
            # --- FLAGS DE VALIDATION OBLIGATOIRES ---
            "mandatory_validation_required": True,
            "validation_notice": (
                "ATTENTION : Ce document DUME/ESPD est un projet pré-rempli à partir des données du dossier d'entreprise. "
                "Tous les champs 'exclusion_grounds' et 'selection_criteria' (valeur null) doivent obligatoirement être "
                "complétés, vérifiés et signés par le représentant légal habilité avant tout dépôt officiel. "
                "Un dépôt avec des valeurs non-vérifiées engage la responsabilité pénale du signataire (Art. L2141-7 CCP)."
            ),
            "declaration_status": "draft_requires_human_validation",
        }


admin_dossier_service = AdminDossierService()
