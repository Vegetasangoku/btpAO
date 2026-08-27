"""
Test Suite for Lot P6:
1. Client Template In-Place Processing (Preserving Styles & Formats).
2. Multi-Tier Source Hierarchy Resolution (Project > RAG > Company Assets > Learnings).
3. Anti-Hallucination [À COMPLÉTER] Red Markers.
4. Completeness Analysis & Checklist Metrics.
"""
import asyncio
import io
import uuid
import docx
import pytest
from fastapi.testclient import TestClient
from jose import jwt
from sqlalchemy import text

from app.core.config import settings
from app.core.db import AsyncSessionLocal
from app.main import app
from app.models.entities import CompanyAsset, Project, Tenant, TenantLearning, User
from app.services.client_template_filler_service import client_template_filler_service

TENANT_ID = "11111111-1111-1111-1111-111111111111"
USER_ID = "33333333-3333-3333-3333-333333333333"
PROJECT_ID = "55555555-5555-5555-5555-555555555555"

JWT_SECRET = settings.SUPABASE_JWT_SECRET or settings.SECRET_KEY
ALGORITHM = "HS256"


def create_token(user_id: str, tenant_id: str) -> str:
    payload = {
        "sub": user_id,
        "email": f"user_{user_id[:8]}@test.fr",
        "tenant_id": tenant_id,
        "app_metadata": {"tenant_id": tenant_id},
        "user_metadata": {"tenant_id": tenant_id},
        "aud": "authenticated",
        "role": "authenticated",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=ALGORITHM)


def create_sample_client_template_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("MÉMOIRE TECHNIQUE — MODÈLE CLIENT EXCLUSIF", level=1)
    doc.add_paragraph("Acheteur Public Adjudicateur : {{client_name}}")
    doc.add_paragraph("Intitulé de l'Opération : {{project_title}}")
    doc.add_paragraph("Numéro SIRET de l'Entreprise : {{siret}}")
    doc.add_paragraph("Agrément Nucléaire Spécifique Inconnu : {{unknown_nuclear_cert_xyz}}")
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Paramètre"
    table.rows[0].cells[1].text = "Valeur"
    table.rows[1].cells[0].text = "Référence Consultation"
    table.rows[1].cells[1].text = "{{reference_code}}"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture(scope="module", autouse=True)
def setup_test_data():
    async def _setup():
        async with AsyncSessionLocal() as db:
            await db.execute(text("DELETE FROM tenant_learnings WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM company_assets WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM generated_sections WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM projects WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM users WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM tenants WHERE id = :t1"), {"t1": TENANT_ID})
            await db.commit()

            t = Tenant(id=uuid.UUID(TENANT_ID), name="BTP Construction Grand Ouest SAS", slug="btp-grand-ouest", siret="98765432100019", country_code="FR")
            db.add(t)
            await db.flush()

            u = User(id=uuid.UUID(USER_ID), tenant_id=uuid.UUID(TENANT_ID), email="u@grandouest.fr", full_name="Directeur Technique", role="admin")
            p = Project(id=uuid.UUID(PROJECT_ID), tenant_id=uuid.UUID(TENANT_ID), title="Construction Médiathèque Haute Performance", reference_code="AO-2026-MED", client_name="Région Bretagne")
            
            # Asset with verified data
            asset = CompanyAsset(
                id=uuid.uuid4(),
                tenant_id=uuid.UUID(TENANT_ID),
                category="profile",
                title="Profil Entreprise",
                metadata_json={"siret": "98765432100019", "insurance": "SMABTP Police 887412"},
                validated_by_user=True,
            )
            
            # Learning
            learning = TenantLearning(
                id=uuid.uuid4(),
                tenant_id=uuid.UUID(TENANT_ID),
                category="prefill_pattern",
                title="Spécification Levage",
                learning_insight="Préférence grue à tour",
                actionable_directive="Utilisation systématique de grues Potain MDT 219 pour le levage.",
                source_outcome="won",
            )

            db.add_all([u, p, asset, learning])
            await db.commit()

    async def _teardown():
        async with AsyncSessionLocal() as db:
            await db.execute(text("DELETE FROM tenant_learnings WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM company_assets WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM generated_sections WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM projects WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM users WHERE tenant_id = :t1"), {"t1": TENANT_ID})
            await db.execute(text("DELETE FROM tenants WHERE id = :t1"), {"t1": TENANT_ID})
            await db.commit()

    asyncio.run(_setup())
    yield
    asyncio.run(_teardown())


def test_client_template_fill_inplace_endpoint():
    client = TestClient(app)
    token = create_token(USER_ID, TENANT_ID)
    template_docx_bytes = create_sample_client_template_docx()

    files = {
        "file": ("template_client.docx", template_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }

    res = client.post(
        f"/api/client-templates/{PROJECT_ID}/fill",
        files=files,
        headers={"Authorization": f"Bearer {token}"},
    )

    assert res.status_code == 200
    assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in res.headers["content-type"]
    assert "X-Completeness-Score" in res.headers
    assert "X-Pending-Actions" in res.headers

    # Validate document content replacement
    doc_filled = docx.Document(io.BytesIO(res.content))
    paragraphs_text = [p.text for p in doc_filled.paragraphs]
    
    # Assert Tier 1 resolution
    assert any("Région Bretagne" in t for t in paragraphs_text)
    assert any("Construction Médiathèque Haute Performance" in t for t in paragraphs_text)
    # Assert Tier 3 resolution
    assert any("98765432100019" in t for t in paragraphs_text)
    # Assert Anti-Hallucination red marker for unknown field
    assert any("[À COMPLÉTER : unknown_nuclear_cert_xyz]" in t for t in paragraphs_text)


def test_client_template_analyze_completeness_endpoint():
    client = TestClient(app)
    token = create_token(USER_ID, TENANT_ID)
    template_docx_bytes = create_sample_client_template_docx()

    files = {
        "file": ("template_client.docx", template_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }

    res = client.post(
        f"/api/client-templates/{PROJECT_ID}/analyze-completeness",
        files=files,
        headers={"Authorization": f"Bearer {token}"},
    )

    assert res.status_code == 200
    data = res.json()
    assert data["total_fields"] == 5 # 4 paragraphs + 1 table cell
    assert data["filled_fields"] == 4
    assert data["pending_actions_count"] == 1
    assert data["completeness_score_pct"] == 80.0
    assert data["is_ready_for_submission"] is False
    assert len(data["sections"]) == 5


def test_client_template_rejects_non_docx_file():
    client = TestClient(app)
    token = create_token(USER_ID, TENANT_ID)

    files = {
        "file": ("document.pdf", b"%PDF-1.4 dummy binary content", "application/pdf")
    }

    res = client.post(
        f"/api/client-templates/{PROJECT_ID}/fill",
        files=files,
        headers={"Authorization": f"Bearer {token}"},
    )
    assert res.status_code == 400
    assert "Seuls les fichiers Word (.docx)" in res.json()["detail"]
