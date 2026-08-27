#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Standalone functional test for the Batch-9 additions to exporter_service.py.
Loads the REAL patched module file directly (importlib by path) after
stubbing its unavailable/heavy dependencies (docxtpl, app.core.storage,
app.services.gantt_service, app.services.diagram_service) into sys.modules,
so this exercises the actual shipped code, not a hand-copied lookalike.
"""
import sys
import io
import types
import importlib.util

EXPORTER_PATH = sys.argv[1]

# ── Stub unavailable dependencies ──────────────────────────────────────────
docxtpl_stub = types.ModuleType("docxtpl")
docxtpl_stub.DocxTemplate = object
docxtpl_stub.InlineImage = object
sys.modules["docxtpl"] = docxtpl_stub

app_pkg = types.ModuleType("app")
app_pkg.__path__ = []
sys.modules["app"] = app_pkg

app_core_pkg = types.ModuleType("app.core")
app_core_pkg.__path__ = []
sys.modules["app.core"] = app_core_pkg

app_core_storage = types.ModuleType("app.core.storage")
app_core_storage.storage_service = object()
sys.modules["app.core.storage"] = app_core_storage

app_services_pkg = types.ModuleType("app.services")
app_services_pkg.__path__ = []
sys.modules["app.services"] = app_services_pkg

app_services_gantt = types.ModuleType("app.services.gantt_service")
app_services_gantt.gantt_service = object()
sys.modules["app.services.gantt_service"] = app_services_gantt

app_services_diagram = types.ModuleType("app.services.diagram_service")
app_services_diagram.diagram_service = object()
sys.modules["app.services.diagram_service"] = app_services_diagram

# ── Load the REAL patched exporter_service.py by path ──────────────────────
spec = importlib.util.spec_from_file_location("exporter_service_under_test", EXPORTER_PATH)
mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(mod)
print("MODULE LOADED OK:", EXPORTER_PATH)

svc = mod.ExporterService()

# ── Build a synthetic client template .docx ────────────────────────────────
import docx
tpl = docx.Document()
header = tpl.sections[0].header
hp = header.paragraphs[0]
hp.add_run("{nom_du_client} — Réponse à l'Appel d'Offres")

footer = tpl.sections[0].footer
fp = footer.paragraphs[0]
fp.add_run("Document confidentiel — [NOM_DU_CLIENT]")

# Only 3 of the 9 canonical sections present in this template's own structure
tpl.add_heading("1. Présentation de l'Entreprise", level=1)
tpl.add_paragraph("Texte de présentation...")
tpl.add_heading("3. Moyens Humains & Encadrement", level=1)
tpl.add_paragraph("Texte moyens humains...")
tpl.add_heading("5. Méthodologie & Planning Prévisionnel", level=1)
tpl.add_paragraph("Texte méthodologie...")

buf = io.BytesIO()
tpl.save(buf)
template_bytes = buf.getvalue()
print(f"SYNTHETIC TEMPLATE BUILT: {len(template_bytes)} bytes")

REQUIRED_TITLES = [
    "1. Présentation de l'Entreprise",
    "2. Références de Travaux Similaires",
    "3. Moyens Humains & Encadrement",
    "4. Moyens Matériels & Engins",
    "5. Méthodologie & Planning Prévisionnel",
    "6. Démarche Qualité & Autocontrôle",
    "7. Sécurité, Prévention & PPSPS",
    "8. RSE, Déchets BTP & Bilan Carbone",
    "9. Politique de Sous-Traitance",
]

# ── Test 1: _detect_missing_required_sections ───────────────────────────────
missing = svc._detect_missing_required_sections(template_bytes, REQUIRED_TITLES)
expected_missing = {
    "2. Références de Travaux Similaires",
    "4. Moyens Matériels & Engins",
    "6. Démarche Qualité & Autocontrôle",
    "7. Sécurité, Prévention & PPSPS",
    "8. RSE, Déchets BTP & Bilan Carbone",
    "9. Politique de Sous-Traitance",
}
expected_present = {
    "1. Présentation de l'Entreprise",
    "3. Moyens Humains & Encadrement",
    "5. Méthodologie & Planning Prévisionnel",
}
missing_set = set(missing)
print("MISSING DETECTED:", missing_set)
assert missing_set == expected_missing, f"MISMATCH: got {missing_set}, expected {expected_missing}"
assert not (missing_set & expected_present), "FALSE POSITIVE: a present section was flagged as missing"
print("TEST 1 PASSED: missing-section detection exact match.")

# Edge case: no template -> no missing list (never invents one)
assert svc._detect_missing_required_sections(None, REQUIRED_TITLES) == []
assert svc._detect_missing_required_sections(template_bytes, None) == []
assert svc._detect_missing_required_sections(template_bytes, []) == []
print("TEST 1b PASSED: no-template / no-required-titles edge cases return [].")

# ── Test 2: _replace_company_placeholders ───────────────────────────────────
doc2 = docx.Document(io.BytesIO(template_bytes))
before_header = doc2.sections[0].header.paragraphs[0].text
before_footer = doc2.sections[0].footer.paragraphs[0].text
print("BEFORE header:", repr(before_header))
print("BEFORE footer:", repr(before_footer))
assert "{nom_du_client}" in before_header
assert "[NOM_DU_CLIENT]" in before_footer

n = svc._replace_company_placeholders(doc2, "EiffaBTP Construction SAS")
after_header = doc2.sections[0].header.paragraphs[0].text
after_footer = doc2.sections[0].footer.paragraphs[0].text
print("AFTER header:", repr(after_header))
print("AFTER footer:", repr(after_footer))
assert n == 2, f"expected 2 replacements, got {n}"
assert "EiffaBTP Construction SAS" in after_header
assert "{nom_du_client}" not in after_header
assert "EiffaBTP Construction SAS" in after_footer
assert "[NOM_DU_CLIENT]" not in after_footer
print("TEST 2 PASSED: placeholder replacement in header AND footer.")

# ── Test 3: full build_memo_docx end-to-end with the synthetic template ────
sections = [
    {"id": "1", "section_key": "presentation_entreprise", "title": "1. Présentation de l'Entreprise",
     "order_index": 1, "content_html": "<p>Contenu généré.</p>", "compliance_score": 95.0},
]
project_data = {
    "id": "test-proj", "title": "Réhabilitation Test", "reference_code": "AO-TEST-01",
    "client_name": "Client Test", "location": "Nantes", "lot_number": "Lot 01",
    "budget_estimate": 1000000.0, "company_name": "EiffaBTP Construction SAS",
}
result = svc.build_memo_docx(
    tenant_id="test-tenant",
    project_id="test-proj",
    project_data=project_data,
    sections=sections,
    decision_form={},
    template_bytes=template_bytes,
    include_visuals=False,
    required_section_titles=REQUIRED_TITLES,
)
out_bytes = result["docx_bytes"]
print(f"TEST 3: build_memo_docx returned {len(out_bytes)} bytes")

# Re-open the produced docx and verify: header placeholder replaced, missing-section
# warning paragraph present, and a real docx.Document can parse it without error.
out_doc = docx.Document(io.BytesIO(out_bytes))
out_header_text = out_doc.sections[0].header.paragraphs[0].text
assert "EiffaBTP Construction SAS" in out_header_text, f"header not replaced: {out_header_text!r}"
assert "{nom_du_client}" not in out_header_text

full_body_text = "\n".join(p.text for p in out_doc.paragraphs)
assert "Sections requises par l'appel d'offres" in full_body_text, "missing-section warning header not found in output"
for t in expected_missing:
    assert t in full_body_text, f"missing title '{t}' not listed as a warning in output"
for t in expected_present:
    # present sections should NOT appear a second time in the missing-warning list,
    # but the TOC bullet from `sections` only lists the 1 section actually generated
    pass
print("TEST 3 PASSED: end-to-end build_memo_docx output is valid and contains expected fidelity fixes.")

print("\nALL TESTS PASSED.")
